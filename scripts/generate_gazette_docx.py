"""Generates gazette-2026-06-20.docx from the Issue No. 2 content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BLACK  = RGBColor(0x1a, 0x1a, 0x1a)
GREY   = RGBColor(0x55, 0x55, 0x55)
LGREY  = RGBColor(0x88, 0x88, 0x88)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
CREAM  = RGBColor(0xFA, 0xFA, 0xFA)
AMBER  = RGBColor(0x6b, 0x59, 0x00)
DKGRN  = RGBColor(0x1a, 0x3a, 0x1a)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_border_bottom(paragraph, color="1a1a1a", size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_border_top(paragraph, color="1a1a1a", size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = paragraph._p.pPr.get_or_add_pPr() if paragraph._p.pPr is not None else OxmlElement("w:pPr")
    bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), str(size))
    top.set(qn("w:space"), "4")
    top.set(qn("w:color"), color)
    bdr.append(top)
    paragraph._p.pPr.append(bdr)

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

def shade_paragraph(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)

def bold_run(para, text, size=None, color=None, font="Georgia"):
    run = para.add_run(text)
    run.bold = True
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def normal_run(para, text, size=None, color=None, italic=False, font="Arial"):
    run = para.add_run(text)
    run.bold = False
    run.italic = italic
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ── Build document ──────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.2)
section.right_margin  = Cm(2.2)
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(2.0)

# ── MASTHEAD ────────────────────────────────────────────────────────────────────

# Top meta line
meta_top = doc.add_paragraph()
meta_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(meta_top, before=0, after=40)
normal_run(meta_top, "ISSN 0000-0000 (PROTOTYPE)   ·   INTERNAL CIRCULATION — OMNI LAW FIRM   ·   omni-law-gazette.internal",
           size=7, color=GREY)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(title_p, before=40, after=20)
r1 = title_p.add_run("OMNI ")
r1.bold = True; r1.font.name = "Georgia"; r1.font.size = Pt(38); r1.font.color.rgb = BLACK
r2 = title_p.add_run("Law ")
r2.bold = False; r2.italic = True; r2.font.name = "Georgia"; r2.font.size = Pt(38); r2.font.color.rgb = BLACK
r3 = title_p.add_run("GAZETTE")
r3.bold = True; r3.font.name = "Georgia"; r3.font.size = Pt(38); r3.font.color.rgb = BLACK

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(sub, before=0, after=40)
normal_run(sub, "LEGISLATIVE INTELLIGENCE FOR LEGAL PROFESSIONALS", size=8.5, color=GREY)

# Thin rule
rule = doc.add_paragraph()
add_border_bottom(rule, color="1a1a1a", size=4)
set_para_spacing(rule, before=0, after=60)

# Issue meta
issue_meta = doc.add_paragraph()
issue_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(issue_meta, before=0, after=80)
normal_run(issue_meta, "ISSUE No. 2     ·     REPORTING PERIOD: 13 JUNE – 20 JUNE 2026     ·     PUBLISHED: FRIDAY, 20 JUNE 2026",
           size=7.5, color=GREY)

# Week banner (simulated with a shaded single-cell table)
banner_tbl = doc.add_table(rows=1, cols=1)
banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
banner_cell = banner_tbl.cell(0, 0)
set_cell_bg(banner_cell, "1a1a1a")
banner_para = banner_cell.paragraphs[0]
banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(banner_para, before=80, after=80)
br = banner_para.add_run("WEEK 25, 2026   ·   REPUBLIC OF AZERBAIJAN   ·   LEGISLATIVE MONITOR")
br.bold = True; br.font.name = "Arial"; br.font.size = Pt(8.5); br.font.color.rgb = WHITE
doc.add_paragraph()  # spacer

# ── AT A GLANCE ──────────────────────────────────────────────────────────────────

glance_hdr = doc.add_paragraph()
set_para_spacing(glance_hdr, before=60, after=40)
normal_run(glance_hdr, "THIS WEEK AT A GLANCE", size=7.5, color=GREY)
add_border_bottom(glance_hdr, color="cccccc", size=4)

stats_tbl = doc.add_table(rows=2, cols=4)
stats_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stats = [
    ("10", "Bills under\nParliamentary Review"),
    ("6",  "Cabinet of Ministers\nDecisions"),
    ("2",  "Cabinet of Ministers\nOrders"),
    ("4",  "Presidential\nDocuments & Addresses"),
]
for i, (num, label) in enumerate(stats):
    cell = stats_tbl.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(num)
    r.bold = True; r.font.name = "Georgia"; r.font.size = Pt(26); r.font.color.rgb = BLACK

    cell2 = stats_tbl.cell(1, i)
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    normal_run(p2, label, size=7.5, color=GREY)

doc.add_paragraph()  # spacer

# ── Helper: Section Header ───────────────────────────────────────────────────────

def section_header(institution, title, subtitle):
    p = doc.add_paragraph()
    set_para_spacing(p, before=160, after=60)
    add_border_bottom(p, color="cccccc", size=4)
    pPr = p._p.get_or_add_pPr()
    bdr_el = OxmlElement("w:pBdr")
    top_el = OxmlElement("w:top")
    top_el.set(qn("w:val"), "single"); top_el.set(qn("w:sz"), "12")
    top_el.set(qn("w:space"), "4"); top_el.set(qn("w:color"), "1a1a1a")
    bdr_el.append(top_el)
    pPr.append(bdr_el)

    r1 = p.add_run(institution.upper() + "   ")
    r1.font.name = "Arial"; r1.font.size = Pt(7); r1.font.color.rgb = LGREY

    r2 = p.add_run(title)
    r2.bold = True; r2.font.name = "Georgia"; r2.font.size = Pt(15); r2.font.color.rgb = BLACK

    r3 = p.add_run("   " + subtitle)
    r3.italic = True; r3.font.name = "Arial"; r3.font.size = Pt(8); r3.font.color.rgb = LGREY

# ── Helper: Doc entry ────────────────────────────────────────────────────────────

def doc_entry(tag, number, date, title, body, tag_color=BLACK):
    # Tag + number + date line
    p1 = doc.add_paragraph()
    set_para_spacing(p1, before=60, after=20)
    rt = p1.add_run(f"[{tag}]")
    rt.bold = True; rt.font.name = "Arial"; rt.font.size = Pt(6.5); rt.font.color.rgb = tag_color
    if number:
        rn = p1.add_run(f"  {number}")
        rn.font.name = "Arial"; rn.font.size = Pt(7.5); rn.font.color.rgb = LGREY
    rd = p1.add_run(f"   {date}")
    rd.font.name = "Arial"; rd.font.size = Pt(7.5); rd.font.color.rgb = LGREY

    # Title
    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=20)
    bold_run(p2, title, size=10, color=BLACK, font="Georgia")

    # Body
    p3 = doc.add_paragraph()
    set_para_spacing(p3, before=0, after=80)
    add_border_bottom(p3, color="d0d0d0", size=2)
    normal_run(p3, body, size=9, color=RGBColor(0x33,0x33,0x33))

def notice_box(text: str):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "f8f8f8")
    p = cell.paragraphs[0]
    set_para_spacing(p, before=60, after=60)
    normal_run(p, text, size=8.5, color=RGBColor(0x55,0x55,0x55))
    doc.add_paragraph()

def spotlight_box(items):
    hdr = doc.add_paragraph()
    set_para_spacing(hdr, before=140, after=60)
    shade_paragraph(hdr, "f0ede6")
    add_border_bottom(hdr, color="cccccc", size=4)
    normal_run(hdr, "REGULATORY SPOTLIGHT — KEY THEMES, WEEK 25 / 2026", size=7.5, color=LGREY)

    for title, body in items:
        p_title = doc.add_paragraph()
        set_para_spacing(p_title, before=60, after=20)
        shade_paragraph(p_title, "f0ede6")
        bold_run(p_title, title, size=9.5, color=BLACK, font="Georgia")

        p_body = doc.add_paragraph()
        set_para_spacing(p_body, before=0, after=60)
        shade_paragraph(p_body, "f0ede6")
        normal_run(p_body, body, size=9, color=RGBColor(0x44,0x44,0x44))

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 1 — MILLI MAJLIS
# ════════════════════════════════════════════════════════════════════════════════

section_header("Parliament of Azerbaijan", "Milli Majlis", "Bills & Legislation — 19 June 2026")

intro = doc.add_paragraph()
set_para_spacing(intro, before=40, after=80)
normal_run(intro, "The Milli Majlis (National Assembly) placed ", size=9, color=GREY)
bold_run(intro, "ten bills", size=9, color=GREY, font="Arial")
normal_run(intro, " on the legislative agenda on ", size=9, color=GREY)
bold_run(intro, "19 June 2026", size=9, color=GREY, font="Arial")
normal_run(intro, " for public discussion and first reading. The package spans fiscal accountability, "
           "trade diplomacy, public administration reform, and a broad multi-sector transport and commerce "
           "omnibus. No new laws or parliamentary decrees were enacted during this reporting period.", size=9, color=GREY)

bills = [
    ("2025 State Budget Execution Report",
     "Annual accountability report on the implementation of the 2025 state budget, submitted for "
     "parliamentary review and approval. A key fiscal governance milestone."),
    ("Customs Code Amendments",
     "Proposed amendments to Azerbaijan's Customs Code covering import/export procedures, tariff "
     "classifications, or enforcement mechanisms. May affect trade and logistics operations."),
    ("Civil Service Law — Amendments",
     "Targeted changes to the Law on State Service addressing civil servant rights, obligations, "
     "conflict-of-interest restrictions, or career advancement frameworks."),
    ("Azerbaijan–Ecuador Visa Exemption Agreement (Diplomatic Passports)",
     "Bilateral visa exemption accord with the Republic of Ecuador for holders of diplomatic passports, "
     "expanding Azerbaijan's diplomatic travel framework in Latin America."),
    ("Constitutional Law on Normative Legal Acts — Amendments",
     "Constitutional-law-level amendments to the framework governing the adoption, publication, and "
     "hierarchy of normative legal acts in Azerbaijan. A structurally significant bill."),
    ("Civil Service & Mass Media Laws — Combined Amendment Package",
     "Dual-law package introducing coordinated changes to state service legislation and the mass media "
     "legal framework. May address public officials' media interactions or media licensing."),
    ("Omnibus Transport & Commerce Amendment Package",
     "A wide-ranging multi-law amendment covering maritime law, road traffic, land transport, customs "
     "duties, civil commerce, automotive regulation, and aviation codes. The most commercially broad "
     "bill of the week, with potential impact on logistics, shipping, and air transport agreements."),
    ("Road Traffic Law Implementation Package",
     "Secondary legislation giving effect to the December 9, 2025 Road Traffic Law reform. Includes "
     "consequential amendments to related acts to align with the new traffic regime."),
    ("Parliament Rules of Procedure & Legislative Initiative Procedures",
     "Reform of Milli Majlis internal regulations and procedures for legislative initiative, potentially "
     "affecting how bills are introduced, debated, and voted upon. Institutional-governance reform."),
    ("General Education Law — Amendments",
     "Proposed changes to the Law on General Education addressing curriculum standards, school "
     "management, teacher certification, or student rights."),
]

for i, (title, desc) in enumerate(bills, 1):
    p = doc.add_paragraph(style="List Number")
    set_para_spacing(p, before=40, after=20)
    bold_run(p, title, size=9.5, color=BLACK, font="Georgia")
    p2 = doc.add_paragraph()
    set_para_spacing(p2, before=0, after=60)
    normal_run(p2, desc, size=8.5, color=GREY)

notice_box("NOTE ON ENACTED LEGISLATION: The most recent laws published in the Official Gazette through "
           "the close of this reporting period remain dated 26 May 2026 (Law No. 415-VIIQD — Amendments "
           "to the Law on Citizens' Appeals). No new laws or parliamentary decrees were enacted during "
           "the week of 13–20 June. The bills listed above are at the discussion stage and have not yet been adopted.")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CABINET OF MINISTERS
# ════════════════════════════════════════════════════════════════════════════════

section_header("İqtisadi Kabinet / Nazirlər Kabineti", "Cabinet of Ministers",
               "Decisions & Orders — 16–19 June 2026")

cab_intro = doc.add_paragraph()
set_para_spacing(cab_intro, before=40, after=80)
normal_run(cab_intro, "The Cabinet of Ministers issued ", size=9, color=GREY)
bold_run(cab_intro, "six Decisions", size=9, color=GREY, font="Arial")
normal_run(cab_intro, " and ", size=9, color=GREY)
bold_run(cab_intro, "two Orders", size=9, color=GREY, font="Arial")
normal_run(cab_intro, " during the reporting period, concentrated on two legislative sessions: a cluster "
           "of three instruments on 16 June addressing digital governance, energy grid regulation, and "
           "civil defence; and a further five instruments on 19 June advancing renewable energy policy "
           "and water resource management.", size=9, color=GREY)

# Decisions sub-header
dh = doc.add_paragraph()
set_para_spacing(dh, before=80, after=40)
add_border_bottom(dh, color="cccccc", size=4)
normal_run(dh, "DECISIONS (QƏRARLAR)", size=8, color=GREY)

decisions = [
    ("No. 181", "16 Jun 2026", "Electricity Grid Operational Guidelines — Amendments",
     "Modifies the existing regulatory framework governing the technical operation of the electricity "
     "transmission and distribution grid. May affect electricity suppliers, large consumers, and grid operators."),
    ("No. 182", "16 Jun 2026", "Electronic State Services Registry — Updated Regulations",
     "Revises regulations governing the official registry of electronic state services and e-service "
     "delivery standards. Advances the government's digital transformation agenda for citizen-facing services."),
    ("No. 183", "16 Jun 2026", "Civil Defence Organisation for Railway Transport",
     "Establishes civil defence organisational procedures applicable to railway transport operations. "
     "Part of a broader effort to integrate critical infrastructure into Azerbaijan's emergency preparedness framework."),
    ("No. 184", "19 Jun 2026", "Charter Documents — Two State Water Management Entities",
     "Approves the founding charter instruments for two new state entities under Azerbaijan's Water "
     "Resources Agency. Establishes institutional structures for water resource management, potentially "
     "linked to irrigated agriculture or the liberated territories."),
    ("No. 185", "19 Jun 2026", "Energy Balance Preparation & Monitoring — Amended Procedures",
     "Amends the procedural rules for preparing and monitoring the national energy balance. Affects "
     "how energy supply and demand data is collected, reported, and used for planning purposes."),
    ("No. 186", "19 Jun 2026", "Solar Power Station Construction — Extended Permission Duration",
     "Extends the validity period of permissions granted for constructing solar power stations within "
     "designated renewable energy source areas, easing timelines for investors in the solar sector."),
]

for num, date, title, body in decisions:
    doc_entry("DECISION", num, date, title, body, tag_color=RGBColor(0x2c,0x3e,0x50))

# Orders sub-header
oh = doc.add_paragraph()
set_para_spacing(oh, before=100, after=40)
add_border_bottom(oh, color="cccccc", size=4)
normal_run(oh, "ORDERS (SƏRƏNCAMLAR)", size=8, color=GREY)

orders = [
    ("No. 419s", "19 Jun 2026", "Airport Management Operational Standards — Enforcement",
     "Directs implementation of airport management operational standards approved by a prior presidential "
     "directive, translating policy into administrative action for airport authorities and relevant ministries."),
    ("No. 420s", "19 Jun 2026", "Citizens' Petitions & Government Office Protocols — Implementation",
     "Operationalises a presidential decree on citizen petition procedures and government office conduct "
     "protocols, establishing implementation measures for responsible agencies. Relevant for administrative law practitioners."),
]

for num, date, title, body in orders:
    doc_entry("ORDER", num, date, title, body, tag_color=RGBColor(0x4a,0x4a,0x4a))

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PRESIDENTIAL OFFICE
# ════════════════════════════════════════════════════════════════════════════════

section_header("Administration of the President", "Presidential Office",
               "Addresses & Diplomatic Correspondence — 13–18 June 2026")

pres_intro = doc.add_paragraph()
set_para_spacing(pres_intro, before=40, after=80)
normal_run(pres_intro,
           "No presidential decrees (fərmanlar) or orders (sərəncamlar) were published in the publicly "
           "accessible document register for the week under review. The presidential record for this period "
           "consists of two substantive addresses at international forums hosted in Azerbaijan, and two "
           "diplomatic letters to foreign heads of state. The two summit addresses — on artificial "
           "intelligence and human rights, and on South Caucasus peacebuilding — are of particular geopolitical significance.",
           size=9, color=GREY)

# Summit addresses
sa_hdr = doc.add_paragraph()
set_para_spacing(sa_hdr, before=60, after=40)
add_border_bottom(sa_hdr, color="cccccc", size=4)
normal_run(sa_hdr, "SUMMIT ADDRESSES", size=8, color=GREY)

doc_entry("ADDRESS", None, "18 Jun 2026",
          "Opening Address — International Baku Ombudsmen Summit",
          "Remarks delivered at the International Baku Ombudsmen Summit on \"Human Rights in the Age of "
          "Artificial Intelligence.\" The summit brought together ombudsmen and human rights institutions "
          "to discuss AI governance, algorithmic accountability, and fundamental rights protection. "
          "Azerbaijan's hosting of this forum reinforces its role as a venue for multilateral human rights dialogue.",
          tag_color=RGBColor(0x5a,0x3a,0x1a))

doc_entry("ADDRESS", None, "15 Jun 2026",
          "Welcome Address — Shusha Peacebuilding & Regional Security Conference",
          "Opening remarks at an international conference on regional security and South Caucasus "
          "peacebuilding, held in Shusha. The choice of Shusha — the cultural capital of Azerbaijan — "
          "as the venue is symbolically significant and underscores Azerbaijan's post-conflict "
          "statebuilding narrative.",
          tag_color=RGBColor(0x5a,0x3a,0x1a))

# Diplomatic letters table
dl_hdr = doc.add_paragraph()
set_para_spacing(dl_hdr, before=100, after=40)
add_border_bottom(dl_hdr, color="cccccc", size=4)
normal_run(dl_hdr, "DIPLOMATIC LETTERS", size=8, color=GREY)

diplo_tbl = doc.add_table(rows=3, cols=5)
diplo_tbl.style = "Table Grid"
diplo_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Date", "Direction", "Counterpart", "Country / Office", "Occasion"]
for j, h in enumerate(headers):
    cell = diplo_tbl.cell(0, j)
    set_cell_bg(cell, "1a1a1a")
    p = cell.paragraphs[0]
    r = p.add_run(h.upper())
    r.bold = True; r.font.name = "Arial"; r.font.size = Pt(7); r.font.color.rgb = WHITE

rows_data = [
    ("14 Jun", "OUTBOUND", "Donald J. Trump",     "President of the United States", "80th Birthday Congratulations"),
    ("13 Jun", "OUTBOUND", "King Charles III",     "United Kingdom",                 "National Holiday (Trooping the Colour)"),
]
for i, row in enumerate(rows_data, 1):
    for j, val in enumerate(row):
        cell = diplo_tbl.cell(i, j)
        p = cell.paragraphs[0]
        if j == 1:
            r = p.add_run(val)
            r.bold = True; r.font.name = "Arial"; r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x4a,0x1a,0x1a)
        else:
            normal_run(p, val, size=8.5)

doc.add_paragraph()

notice_box("NOTE ON DECREES & ORDERS: Presidential decrees and orders were not retrievable from the "
           "president.az documents portal for the current reporting period. Readers should verify "
           "directly via the president.az documents register for any decrees or orders issued between "
           "13–20 June. Cabinet Orders 419s and 420s reference implementing prior presidential directives, "
           "confirming regulatory activity at the executive level.")

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 4 — CONSTITUTIONAL COURT
# ════════════════════════════════════════════════════════════════════════════════

section_header("Ali Məhkəmə Orqanı", "Constitutional Court",
               "Decisions — constcourt.gov.az")

ct_tbl = doc.add_table(rows=1, cols=1)
ct_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
ct_cell = ct_tbl.cell(0, 0)
set_cell_bg(ct_cell, "fff8e1")
ct_p = ct_cell.paragraphs[0]
set_para_spacing(ct_p, before=60, after=60)
bold_run(ct_p, "⚠ Source Temporarily Unavailable.  ", size=8.5, color=AMBER, font="Arial")
normal_run(ct_p,
           "The Constitutional Court's decisions portal (constcourt.gov.az) returned an HTTP 503 "
           "Service Unavailable error at the time of compilation (20 June 2026). No decisions from "
           "the reporting period could be retrieved. Readers are advised to consult the portal directly "
           "for any decisions issued during the week of 13–20 June 2026.",
           size=8.5, color=AMBER)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
# SECTION 5 — E-QANUN
# ════════════════════════════════════════════════════════════════════════════════

section_header("Official Legislative Database", "e-Qanun",
               "Normative Acts Register — e-qanun.az")

eq_tbl = doc.add_table(rows=1, cols=1)
eq_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
eq_cell = eq_tbl.cell(0, 0)
set_cell_bg(eq_cell, "fff8e1")
eq_p = eq_cell.paragraphs[0]
set_para_spacing(eq_p, before=60, after=60)
bold_run(eq_p, "⚠ Data Not Extractable.  ", size=8.5, color=AMBER, font="Arial")
normal_run(eq_p,
           "The e-Qanun portal (e-qanun.az) renders its content dynamically via JavaScript and does "
           "not expose a static listing accessible to automated retrieval. No enacted normative acts "
           "could be extracted for automated review this week. Readers should consult the portal "
           "directly for the consolidated texts of any legislation enacted or amended during the reporting period.",
           size=8.5, color=AMBER)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════════
# REGULATORY SPOTLIGHT
# ════════════════════════════════════════════════════════════════════════════════

spotlight_box([
    ("1. Energy Sector: Solar & Grid Reform",
     "Three Cabinet instruments this week directly target the energy sector. Decision 186 extends solar "
     "construction permit durations — a concrete incentive for renewable project developers facing timeline "
     "pressure. Decision 185 updates energy balance monitoring, improving planning data quality. Decision 181 "
     "amends electricity grid operational rules, with potential implications for network access and tariff "
     "arrangements. Taken together, these represent a sustained regulatory push aligned with Azerbaijan's "
     "post-COP29 green energy commitments."),
    ("2. Transport & Commerce Omnibus Bill",
     "The Milli Majlis placed an exceptionally broad transport and commerce amendment package on the agenda "
     "(Bill 7), spanning maritime, road, customs, automotive, and aviation law simultaneously. If adopted in "
     "its current form, this omnibus would represent one of the most extensive multi-sector regulatory updates "
     "in years. Firms with clients in shipping, freight, civil aviation, or cross-border trade should track "
     "this bill closely through subsequent readings."),
    ("3. AI & Human Rights Governance",
     "The International Baku Ombudsmen Summit on \"Human Rights in the Age of Artificial Intelligence\" "
     "(18 June) positions Azerbaijan as an active participant in emerging global AI governance frameworks. "
     "While not yet legislative, Summit outcomes may inform future domestic AI-related regulation. Legal "
     "practitioners advising tech companies, data processors, or public-sector bodies on compliance should "
     "monitor any declarations or recommendations arising from this event."),
    ("4. Normative Acts Constitutional Reform",
     "Bill 5 — a constitutional-law-level amendment to the framework governing normative legal acts — is "
     "the most structurally significant legislation before Parliament this week. Changes to the hierarchy, "
     "adoption, and publication rules for normative acts can have downstream effects on the entire regulatory "
     "landscape. Practitioners across all practice areas should watch for the bill's text and subsequent "
     "readings, as it may alter how regulations are interpreted and challenged."),
])

# ════════════════════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════════════════════

footer_rule = doc.add_paragraph()
add_border_bottom(footer_rule, color="1a1a1a", size=6)
set_para_spacing(footer_rule, before=160, after=60)

footer_tbl = doc.add_table(rows=1, cols=2)
footer_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

fl = footer_tbl.cell(0, 0)
fp = fl.paragraphs[0]
set_para_spacing(fp, before=0, after=0)
bold_run(fp, "Omni Law Gazette", size=7.5, color=GREY, font="Arial")
normal_run(fp, " is an internal digest compiled by Omni Law Firm for informational purposes only. "
           "It does not constitute legal advice. Document descriptions are derived from official "
           "government portals; readers should consult primary sources and full texts before relying "
           "on any item. Sources: president.az · nk.gov.az · meclis.gov.az · e-qanun.az · constcourt.gov.az",
           size=7.5, color=GREY)

fr = footer_tbl.cell(0, 1)
fp2 = fr.paragraphs[0]
fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_spacing(fp2, before=0, after=0)
bold_run(fp2, "Issue No. 2\n", size=7.5, color=GREY, font="Arial")
normal_run(fp2, "Week of 13–20 June 2026\nCompiled: 20 June 2026\n© 2026 Omni Law Firm", size=7.5, color=GREY)

# ── Save ─────────────────────────────────────────────────────────────────────────

out = "/home/user/omnilawgazette/scripts/gazettes/gazette-2026-06-20.docx"
doc.save(out)
print(f"Saved: {out}")
