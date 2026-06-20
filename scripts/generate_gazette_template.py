#!/usr/bin/env python3
"""Generate the Omni Law Gazette in the bilingual numbered-digest template
(matching OmniLawGazetteQC...EN/AZ.docx), for the reporting period 13-20 June 2026.

Produces two files:
  scripts/gazettes/OmniLawGazette20260613to20EN.docx
  scripts/gazettes/OmniLawGazette20260613to20AZ.docx
"""
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette (from template) ──
NAVY  = RGBColor(0x1F, 0x38, 0x64)
GREY  = RGBColor(0x3B, 0x38, 0x38)
OLIVE = RGBColor(0x80, 0x60, 0x00)
GOLD  = RGBColor(0x8A, 0x6D, 0x1E)
BLACK = RGBColor(0x00, 0x00, 0x00)

SERIF  = "Times New Roman"
SANS   = "Eloquia Text ExtLt"

# Page geometry (A4, 0.5in margins) — EMU values taken from the template
PAGE_W, PAGE_H = 7562215, 10689590
MARGIN = 457200
TEXT_W = PAGE_W - 2 * MARGIN


def _spacing(p, before=0, after=0, line=240):
    """before/after in points, line in 240ths (240 = single, 252 = 1.05)."""
    pPr = p._p.get_or_add_pPr()
    el = OxmlElement("w:spacing")
    el.set(qn("w:before"), str(int(before * 20)))
    el.set(qn("w:after"), str(int(after * 20)))
    el.set(qn("w:line"), str(line))
    el.set(qn("w:lineRule"), "auto")
    pPr.append(el)


def _bottom_border(p, sz=4, color="1F3864"):
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    bdr.append(bottom)
    pPr.append(bdr)


def _run(p, text, font=SANS, size=9, bold=False, italic=False, color=BLACK):
    r = p.add_run(text)
    r.font.name = font
    # ensure complex/East-European script also uses the same font
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), font)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r


def build(lang: str, data: dict, out_path: str):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Emu(PAGE_W), Emu(PAGE_H)
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Emu(MARGIN)

    # ── Masthead title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _spacing(p, 0, 0, 240)
    _bottom_border(p, sz=6, color="1F3864")
    _run(p, "omni", font=SANS, size=54, color=BLACK)
    _run(p, "  law ", font=SANS, size=17, color=BLACK)
    _run(p, "gazette", font=SANS, size=17, color=NAVY)

    # ── Subtitle line with right tab ──
    p = doc.add_paragraph()
    _spacing(p, 0, 8, 240)
    _bottom_border(p, sz=4, color="1F3864")
    p.paragraph_format.tab_stops.add_tab_stop(Emu(TEXT_W), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, data["sub_left"], font=SANS, size=9, bold=True, color=NAVY)
    _run(p, "\t", font=SANS, size=9)
    _run(p, data["sub_right"], font=SANS, size=9, color=GREY)

    # spacer
    _spacing(doc.add_paragraph(), 0, 0, 240)

    # ── Items ──
    for i, it in enumerate(data["items"], 1):
        # number
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, 8, 0, 240)
        _run(p, f"{i:02d}", font=SANS, size=36, bold=True, color=GREY)
        # category
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, 0, 1, 240)
        _run(p, it["cat"], font=SERIF, size=7, bold=True, color=OLIVE)
        # headline
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, 0, 4, 240)
        _run(p, it["head"], font=SERIF, size=13, bold=True, color=NAVY)
        # body paragraphs
        for body in it["body"]:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _spacing(p, 0, 4, 252)
            _run(p, body, font=SANS, size=9, color=BLACK)
        # source
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(p, 0, 6, 240)
        _run(p, it["src"], font=SERIF, size=8, italic=True, color=GOLD)

    # ── Footer rule + disclaimer ──
    p = doc.add_paragraph()
    _spacing(p, 8, 0, 240)
    _bottom_border(p, sz=4, color="1F3864")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(p, 4, 2, 240)
    _run(p, "Omni Law Gazette", font=SANS, size=7.5, bold=True, color=NAVY)
    _run(p, data["foot1"], font=SANS, size=7.5, color=GREY)

    p = doc.add_paragraph()
    _spacing(p, 0, 0, 240)
    _run(p, data["foot2"], font=SANS, size=7.5, bold=True, color=GREY)

    doc.save(out_path)
    print("Saved:", out_path)


# ══════════════════════════════════════════════════════════════════════════
# CONTENT — 13–20 June 2026
# ══════════════════════════════════════════════════════════════════════════

EN = {
    "sub_left": "WEEKLY LEGISLATIVE DIGEST · REPUBLIC OF AZERBAIJAN",
    "sub_right": "WEEKLY EDITION · 13–20 JUNE 2026",
    "foot1": ("  ·  Internal weekly digest of the legislative activity of the Republic of "
              "Azerbaijan, compiled by Omni Law Firm for informational purposes only. It does "
              "not constitute legal advice; consult the consolidated texts on e-qanun.az before "
              "relying on any item. Sources: president.az · nk.gov.az · meclis.gov.az · "
              "e-qanun.az · constcourt.gov.az"),
    "foot2": "Weekly edition  ·  Reporting period 13–20 June 2026  ·  © 2026 Omni Law Firm",
    "items": [
        {
            "cat": "PRESIDENT · DECREE · 16 JUNE 2026",
            "head": "Additional measures are adopted to support non-oil and gas exports.",
            "body": [
                "By a decree of 16 June, the President approved additional measures to support the "
                "export of non-oil-and-gas goods. The instrument forms part of the State's policy of "
                "economic diversification and of reducing the economy's dependence on hydrocarbon revenues.",
                "The decree is of relevance to exporters and to businesses operating in the non-oil "
                "sector, which should review the support mechanisms and any eligibility conditions "
                "established under it once the implementing rules are published.",
            ],
            "src": "Source · e-qanun.az/framework/62008",
        },
        {
            "cat": "PRESIDENT · ORDER · 17 JUNE 2026",
            "head": "The “Dostluq” Order is conferred on Mohammed Suleyman Al-Jasir.",
            "body": [
                "By a presidential order of 17 June, Mr Mohammed Suleyman Al-Jasir was awarded the "
                "“Dostluq” (Friendship) Order in recognition of his contribution to the development of "
                "cooperation with the Republic of Azerbaijan.",
                "State decorations of this kind mark contributions to bilateral and multilateral "
                "cooperation; the award is of note for its diplomatic significance rather than for any "
                "direct legal effect.",
            ],
            "src": "Source · e-qanun.az/framework/62010",
        },
        {
            "cat": "PRESIDENT · ORDER · 17 JUNE 2026",
            "head": "State decorations are conferred on healthcare workers.",
            "body": [
                "By a presidential order of 17 June, a group of healthcare workers of the Republic of "
                "Azerbaijan were conferred State awards in recognition of their services in the "
                "protection of public health.",
                "The order is a customary recognition of the medical profession and carries no general "
                "regulatory effect.",
            ],
            "src": "Source · e-qanun.az/framework/62009",
        },
        {
            "cat": "CABINET · DECISION No. 186 · 19 JUNE 2026",
            "head": "The permit period for building solar power stations is extended.",
            "body": [
                "The Cabinet of Ministers adopted a decision extending the period for obtaining a "
                "permit for the construction of a solar power station within renewable-energy-source areas.",
                "The change eases project timelines for solar developers and is consistent with the "
                "State's renewable-energy targets. Investors in the sector should take account of the "
                "revised permit deadlines.",
            ],
            "src": "Source · nk.gov.az · Decision No. 186",
        },
        {
            "cat": "CABINET · DECISION No. 185 · 19 JUNE 2026",
            "head": "The rules on the preparation of the energy balance are amended.",
            "body": [
                "The Cabinet of Ministers amended the Rules on the preparation, approval and monitoring "
                "of the energy balance, approved by its Decision No. 233 of 1 May 2024.",
                "The amendment affects how national energy supply-and-demand data is compiled and "
                "supervised, and is relevant to bodies engaged in energy planning and reporting.",
            ],
            "src": "Source · nk.gov.az · Decision No. 185",
        },
        {
            "cat": "CABINET · DECISION No. 184 · 19 JUNE 2026",
            "head": "Charters of two State water-management companies are approved.",
            "body": [
                "The Cabinet of Ministers approved the charters of the “Regional Water Melioration” and "
                "“Large Cities Water Supply” closed joint-stock companies, which operate under the "
                "Azerbaijan State Water Resources Agency.",
                "The instrument establishes the legal and institutional framework for these utilities and "
                "may bear on water-infrastructure and irrigation projects, including in the restored territories.",
            ],
            "src": "Source · nk.gov.az · Decision No. 184",
        },
        {
            "cat": "CABINET · DECISION No. 183 · 16 JUNE 2026",
            "head": "Rules are adopted for civil defence in railway transport.",
            "body": [
                "The Cabinet of Ministers approved the Rules on the organisation of civil defence in "
                "railway transport. The instrument integrates rail operations into the State's "
                "emergency-preparedness framework.",
                "Railway operators and related undertakings should align their internal procedures with "
                "the new civil-defence requirements.",
            ],
            "src": "Source · e-qanun.az/framework/62026",
        },
        {
            "cat": "CABINET · DECISION No. 182 · 16 JUNE 2026",
            "head": "The rules on the Electronic Registry of State Services are revised.",
            "body": [
                "The Cabinet of Ministers amended the Rules for maintaining the Electronic Registry of "
                "State Services (approved by Decision No. 32 of 9 February 2015) and the rules on the "
                "organisation and provision of electronic services (approved by Decision No. 241 of 4 August 2025).",
                "The amendments advance the digitalisation of citizen-facing public services and are "
                "relevant to bodies that deliver State e-services.",
            ],
            "src": "Source · e-qanun.az/framework/62025",
        },
        {
            "cat": "CABINET · DECISION No. 181 · 16 JUNE 2026",
            "head": "The Rules on the Electricity Grid are amended.",
            "body": [
                "The Cabinet of Ministers amended the “Rules on the Electricity Grid”, approved by its "
                "Decision No. 315 of 25 June 2024.",
                "The amendment concerns the technical operation of the transmission and distribution "
                "network and may affect electricity suppliers, large consumers and grid operators.",
            ],
            "src": "Source · e-qanun.az/framework/62027",
        },
        {
            "cat": "CABINET · ORDER No. 419s · 19 JUNE 2026",
            "head": "The new airport-management rules are brought into effect.",
            "body": [
                "The Cabinet of Ministers issued an order giving effect to Presidential Decree No. 684 "
                "of 12 June 2026, which approved the Rules on the management of airports.",
                "The order operationalises the airport-management framework and is of relevance to "
                "airport operators and to the civil-aviation authorities.",
            ],
            "src": "Source · nk.gov.az · Order No. 419s",
        },
        {
            "cat": "CABINET · ORDER No. 420s · 19 JUNE 2026",
            "head": "Implementing measures are issued for the Citizens' Appeals Law amendments.",
            "body": [
                "The Cabinet of Ministers issued an order giving effect to Law No. 415-VIIQD of 26 May "
                "2026, which amended the Law on Citizens' Appeals.",
                "The order establishes the administrative measures required to apply the amended appeals "
                "regime across government offices, and is relevant to public-administration and "
                "administrative-law practice.",
            ],
            "src": "Source · nk.gov.az · Order No. 420s",
        },
        {
            "cat": "MILLI MAJLIS · BILLS · 19 JUNE 2026",
            "head": "Ten bills are placed on the parliamentary agenda.",
            "body": [
                "On 19 June the Milli Majlis tabled ten bills for discussion, including the 2025 State "
                "Budget execution report, amendments to the Customs Code and the Civil Service Law, a "
                "visa-exemption agreement with Ecuador for diplomatic-passport holders, a "
                "constitutional-law amendment on normative legal acts, a broad omnibus touching maritime, "
                "road, aviation and customs law, and amendments to the General Education Law.",
                "These bills are at the discussion stage and have not been adopted. No new laws, "
                "parliamentary decrees or Constitutional Court decisions were enacted in the review period "
                "— the most recent law remains No. 415-VIIQD of 26 May 2026.",
            ],
            "src": "Source · meclis.gov.az",
        },
    ],
}

AZ = {
    "sub_left": "HƏFTƏLİK QANUNVERİCİLİK İCMALI · AZƏRBAYCAN RESPUBLİKASI",
    "sub_right": "HƏFTƏLİK BURAXILIŞ · 13–20 İYUN 2026",
    "foot1": ("  ·  Azərbaycan Respublikasının qanunvericilik fəaliyyətinin həftəlik daxili icmalı; "
              "Omni Law Firm tərəfindən yalnız məlumat məqsədilə hazırlanıb. Hüquqi məsləhət təşkil "
              "etmir; istənilən maddəyə istinad etməzdən əvvəl e-qanun.az saytındakı konsolidə edilmiş "
              "mətnlərə baxın. Mənbələr: president.az · nk.gov.az · meclis.gov.az · e-qanun.az · constcourt.gov.az"),
    "foot2": "Həftəlik buraxılış  ·  Baxış dövrü 13–20 iyun 2026  ·  © 2026 Omni Law Firm",
    "items": [
        {
            "cat": "PREZİDENT · FƏRMAN · 16 İYUN 2026",
            "head": "Qeyri-neft-qaz ixracının dəstəklənməsi üçün əlavə tədbirlər müəyyən edilir.",
            "body": [
                "Prezident 16 iyun tarixli fərmanı ilə qeyri-neft-qaz malları üzrə ixracın dəstəklənməsi "
                "ilə bağlı əlavə tədbirləri təsdiq etmişdir. Sənəd iqtisadiyyatın şaxələndirilməsi və "
                "karbohidrogen gəlirlərindən asılılığın azaldılması siyasətinin tərkib hissəsidir.",
                "Fərman ixracatçılar və qeyri-neft sektorunda fəaliyyət göstərən sahibkarlar üçün "
                "əhəmiyyət daşıyır; icra qaydaları dərc edildikdən sonra onlar müəyyən edilmiş dəstək "
                "mexanizmlərini və müvafiq şərtləri nəzərdən keçirməlidirlər.",
            ],
            "src": "Mənbə · e-qanun.az/framework/62008",
        },
        {
            "cat": "PREZİDENT · SƏRƏNCAM · 17 İYUN 2026",
            "head": "Məhəmməd Süleyman Əl-Casirə “Dostluq” ordeni verilir.",
            "body": [
                "Prezidentin 17 iyun tarixli sərəncamı ilə Məhəmməd Süleyman Əl-Casir Azərbaycan "
                "Respublikası ilə əməkdaşlığın inkişafındakı töhfəsinə görə “Dostluq” ordeni ilə təltif "
                "edilmişdir.",
                "Bu cür dövlət təltifləri ikitərəfli və çoxtərəfli əməkdaşlığa verilən töhfələri qeyd "
                "edir; təltif birbaşa hüquqi nəticəsindən çox diplomatik əhəmiyyəti ilə diqqət çəkir.",
            ],
            "src": "Mənbə · e-qanun.az/framework/62010",
        },
        {
            "cat": "PREZİDENT · SƏRƏNCAM · 17 İYUN 2026",
            "head": "Səhiyyə işçiləri dövlət təltiflərinə layiq görülür.",
            "body": [
                "Prezidentin 17 iyun tarixli sərəncamı ilə Azərbaycan Respublikasının bir qrup səhiyyə "
                "işçisi əhalinin sağlamlığının qorunması sahəsindəki xidmətlərinə görə dövlət "
                "təltiflərinə layiq görülmüşdür.",
                "Sərəncam tibb peşəsinə ənənəvi ehtiramın ifadəsidir və ümumi tənzimləyici təsir daşımır.",
            ],
            "src": "Mənbə · e-qanun.az/framework/62009",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · QƏRAR № 186 · 19 İYUN 2026",
            "head": "Günəş elektrik stansiyalarının tikintisi üçün icazə müddəti uzadılır.",
            "body": [
                "Nazirlər Kabineti bərpa olunan enerji mənbələrinin ərazisində günəş elektrik "
                "stansiyasının tikintisinə icazə alınması müddətinin uzadılması haqqında qərar qəbul etmişdir.",
                "Dəyişiklik günəş enerjisi layihələrinin icra müddətlərini yüngülləşdirir və dövlətin "
                "bərpa olunan enerji hədəfləri ilə uzlaşır. Sahəyə investisiya yatıran şəxslər yenilənmiş "
                "icazə müddətlərini nəzərə almalıdırlar.",
            ],
            "src": "Mənbə · nk.gov.az · Qərar № 186",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · QƏRAR № 185 · 19 İYUN 2026",
            "head": "Enerji balansının hazırlanması qaydalarına dəyişiklik edilir.",
            "body": [
                "Nazirlər Kabineti 2024-cü il 1 may tarixli 233 nömrəli Qərarı ilə təsdiq edilmiş "
                "“Enerji balansının hazırlanması, təsdiq edilməsi və icrasına nəzarət Qaydası”na "
                "dəyişiklik etmişdir.",
                "Dəyişiklik milli enerji tələb-təklif məlumatlarının necə hazırlanması və ona nəzarət "
                "olunmasına təsir göstərir; enerji planlaması və hesabatı ilə məşğul olan qurumlar üçün "
                "əhəmiyyət daşıyır.",
            ],
            "src": "Mənbə · nk.gov.az · Qərar № 185",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · QƏRAR № 184 · 19 İYUN 2026",
            "head": "İki dövlət su təsərrüfatı şirkətinin nizamnaməsi təsdiqlənir.",
            "body": [
                "Nazirlər Kabineti Azərbaycan Dövlət Su Ehtiyatları Agentliyinin tabeliyində olan "
                "“Regional Su Meliorasiya” və “İri Şəhərlərin Su Təchizatı” qapalı səhmdar cəmiyyətlərinin "
                "nizamnamələrini təsdiq etmişdir.",
                "Sənəd bu kommunal müəssisələrin hüquqi və institusional çərçivəsini müəyyən edir və su "
                "infrastrukturu və meliorasiya layihələrinə, o cümlədən bərpa olunan ərazilərdə, təsir "
                "göstərə bilər.",
            ],
            "src": "Mənbə · nk.gov.az · Qərar № 184",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · QƏRAR № 183 · 16 İYUN 2026",
            "head": "Dəmir yolu nəqliyyatında mülki müdafiənin təşkili qaydaları təsdiqlənir.",
            "body": [
                "Nazirlər Kabineti “Dəmiryol nəqliyyatında mülki müdafiənin təşkili Qaydaları”nı təsdiq "
                "etmişdir. Sənəd dəmir yolu fəaliyyətini dövlətin fövqəladə hallara hazırlıq çərçivəsinə "
                "daxil edir.",
                "Dəmir yolu operatorları və əlaqəli müəssisələr daxili prosedurlarını yeni mülki müdafiə "
                "tələbləri ilə uzlaşdırmalıdırlar.",
            ],
            "src": "Mənbə · e-qanun.az/framework/62026",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · QƏRAR № 182 · 16 İYUN 2026",
            "head": "Dövlət Xidmətlərinin Elektron Reyestri qaydaları yenilənir.",
            "body": [
                "Nazirlər Kabineti “Dövlət Xidmətlərinin Elektron Reyestrinin aparılması Qaydaları”na "
                "(2015-ci il 9 fevral tarixli 32 nömrəli Qərar) və “Elektron xidmətlərin təşkili və "
                "göstərilməsi qaydaları”na (2025-ci il 4 avqust tarixli 241 nömrəli Qərar) dəyişiklik etmişdir.",
                "Dəyişikliklər vətəndaşlara yönəlmiş dövlət xidmətlərinin rəqəmsallaşmasını irəli aparır "
                "və dövlət elektron xidmətləri göstərən qurumlar üçün əhəmiyyət daşıyır.",
            ],
            "src": "Mənbə · e-qanun.az/framework/62025",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · QƏRAR № 181 · 16 İYUN 2026",
            "head": "“Elektrik şəbəkəsinə dair Qaydalar”a dəyişiklik edilir.",
            "body": [
                "Nazirlər Kabineti 2024-cü il 25 iyun tarixli 315 nömrəli Qərarı ilə təsdiq edilmiş "
                "“Elektrik şəbəkəsinə dair Qaydalar”da dəyişiklik etmişdir.",
                "Dəyişiklik ötürücü və paylayıcı şəbəkənin texniki istismarına aiddir və elektrik "
                "enerjisi təchizatçılarına, iri istehlakçılara və şəbəkə operatorlarına təsir göstərə bilər.",
            ],
            "src": "Mənbə · e-qanun.az/framework/62027",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · SƏRƏNCAM № 419s · 19 İYUN 2026",
            "head": "Hava limanlarının idarə edilməsi qaydaları qüvvəyə minir.",
            "body": [
                "Nazirlər Kabineti “Hava limanlarının (aeroportlarının) idarə edilməsi Qaydası”nın təsdiq "
                "edilməsinə dair Prezidentin 2026-cı il 12 iyun tarixli 684 nömrəli Fərmanının icrasını "
                "təmin edən sərəncam vermişdir.",
                "Sərəncam hava limanlarının idarəetmə çərçivəsini işlək hala gətirir və hava limanı "
                "operatorları və mülki aviasiya orqanları üçün əhəmiyyət daşıyır.",
            ],
            "src": "Mənbə · nk.gov.az · Sərəncam № 419s",
        },
        {
            "cat": "NAZİRLƏR KABİNETİ · SƏRƏNCAM № 420s · 19 İYUN 2026",
            "head": "“Vətəndaşların müraciətləri” qanununa dəyişikliklərin tətbiqi təmin edilir.",
            "body": [
                "Nazirlər Kabineti “Vətəndaşların müraciətləri haqqında” Qanunda dəyişiklik edilməsinə "
                "dair 2026-cı il 26 may tarixli 415-VIIQD nömrəli Qanunun icrasını təmin edən sərəncam vermişdir.",
                "Sərəncam dəyişdirilmiş müraciət rejiminin dövlət qurumlarında tətbiqi üçün zəruri inzibati "
                "tədbirləri müəyyən edir və dövlət idarəçiliyi və inzibati hüquq praktikası üçün əhəmiyyət daşıyır.",
            ],
            "src": "Mənbə · nk.gov.az · Sərəncam № 420s",
        },
        {
            "cat": "MİLLİ MƏCLİS · QANUN LAYİHƏLƏRİ · 19 İYUN 2026",
            "head": "On qanun layihəsi parlament gündəliyinə çıxarılır.",
            "body": [
                "İyunun 19-da Milli Məclis müzakirə üçün on qanun layihəsini gündəliyə çıxarmışdır; "
                "bunların arasında 2025-ci il dövlət büdcəsinin icrası haqqında hesabat, Gömrük Məcəlləsinə "
                "və “Dövlət qulluğu haqqında” Qanuna dəyişikliklər, Ekvadorla diplomatik pasport sahibləri "
                "üçün vizasız rejim haqqında saziş, normativ hüquqi aktlara dair konstitusiya qanununa "
                "düzəliş, dəniz, yol, aviasiya və gömrük hüququnu əhatə edən geniş omnibus paketi, habelə "
                "“Ümumi təhsil haqqında” Qanuna dəyişikliklər vardır.",
                "Bu layihələr müzakirə mərhələsindədir və hələ qəbul edilməmişdir. Baxış dövründə yeni "
                "qanun, parlament qərarı və ya Konstitusiya Məhkəməsinin qərarı qəbul edilməmişdir — ən son "
                "qanun 2026-cı il 26 may tarixli 415-VIIQD nömrəli Qanun olaraq qalır.",
            ],
            "src": "Mənbə · meclis.gov.az",
        },
    ],
}


if __name__ == "__main__":
    import os
    outdir = os.path.join(os.path.dirname(__file__), "gazettes")
    os.makedirs(outdir, exist_ok=True)
    build("EN", EN, os.path.join(outdir, "OmniLawGazette20260613to20EN.docx"))
    build("AZ", AZ, os.path.join(outdir, "OmniLawGazette20260613to20AZ.docx"))
