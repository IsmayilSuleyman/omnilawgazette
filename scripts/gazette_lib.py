#!/usr/bin/env python3
"""Rendering engine for the Omni Law Gazette (house newspaper-digest format).

Shared by the per-issue build scripts. Produces, for a given language, a native
editable .docx and a matching .html twin (the .html is rendered to .pdf via
Chromium by the build scripts). Layout/typography follow the firm's example
edition; see AGENTS.md for the editorial rules.

Data contract — each issue supplies:
  chrome:   {'AZ': {...}, 'EN': {...}}   masthead banner / issue label / footer
  lead:     one article dict (item 01)
  articles: list of article dicts (items 02..N)
where an article dict is:
  {'n': '02',
   'kicker':   {'AZ': str, 'EN': str},
   'headline': {'AZ': str, 'EN': str},
   'body':     {'AZ': [str, ...], 'EN': [str, ...]},
   'source':   {'AZ': str, 'EN': str}}

AUTHORING ORDER: Azerbaijani ('AZ') is the source-of-truth, authored first
and faithful to the official portal wording; English ('EN') is its translation.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import html as _html

NAVY  = RGBColor(0x1F, 0x38, 0x64)
GREY  = RGBColor(0x3B, 0x38, 0x38)
GOLD  = RGBColor(0x80, 0x60, 0x00)
GOLD2 = RGBColor(0x8A, 0x6D, 0x1E)
BLACK = RGBColor(0x00, 0x00, 0x00)
BODY_FONT = 'Eloquia Text ExtLt'
SER_FONT  = 'Times New Roman'
LANGS = ['AZ', 'EN']  # Azerbaijani primary, English translation


# ── DOCX helpers ─────────────────────────────────────────────────────────────
def _set_cols(section, num):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), '708')


def _run(p, text, *, font=BODY_FONT, size=9, bold=False, italic=False, color=None):
    r = p.add_run(text)
    rPr = r._r.get_or_add_rPr(); rfonts = rPr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rPr.insert(0, rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), font)
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def _para(doc, align=None, before=0, after=0, line=1.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    return p


def _bottom_border(el, color='1F3864', sz='4'):
    pPr = el._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '4'); b.set(qn('w:color'), color)
    pbdr.append(b)


def build_docx(lang, basename, chrome, lead, articles, outdir):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = BODY_FONT; st.font.size = Pt(9)
    st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.0
    ch = chrome[lang]

    # masthead (brand "omni law gazette" in both editions)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, after=0)
    _run(p, 'omni', size=54, color=BLACK)
    _run(p, '  law ', size=17, color=BLACK)
    _run(p, 'gazette', size=17, color=NAVY)
    _bottom_border(p, sz='6')

    # banner
    p = _para(doc, after=8)
    _run(p, ch['banner'], size=9, bold=True, color=NAVY)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.27), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, '\t', size=9)
    _run(p, ch['issue'], size=9, color=GREY)
    _bottom_border(p, sz='4')
    _para(doc, after=4)

    def render(item, lead=False):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=(8 if lead else 12), after=0)
        _run(p, item['n'], size=(36 if lead else 28), bold=True, color=GREY)
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=1)
        _run(p, item['kicker'][lang], font=SER_FONT, size=7, bold=True, color=GOLD)
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
        _run(p, item['headline'][lang], font=SER_FONT, size=(13 if lead else 12), bold=True, color=NAVY)
        for par in item['body'][lang]:
            p = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4, line=1.05)
            _run(p, par, size=9, color=BLACK)
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6)
        _run(p, item['source'][lang], font=SER_FONT, size=8, italic=True, color=GOLD2)

    render(lead, lead=True)
    doc.add_section(WD_SECTION.CONTINUOUS); _set_cols(doc.sections[0], 1)
    for item in articles:
        render(item)
    doc.add_section(WD_SECTION.CONTINUOUS); _set_cols(doc.sections[1], 3)

    p = _para(doc, before=6, after=2); _bottom_border(p, sz='4')
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=2)
    _run(p, ch['footer_lead'], size=7.5, bold=True, color=NAVY)
    _run(p, ch['footer'], size=7.5, color=GREY)
    p = _para(doc, after=0)
    _run(p, ch['footer_c'], size=7.5, bold=True, color=GREY)
    _set_cols(doc.sections[2], 1)

    for s in doc.sections:
        s.page_width = Inches(8.27); s.page_height = Inches(11.69)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.5)

    out = os.path.join(outdir, f'{basename}-{lang}.docx')
    doc.save(out)
    return out


# ── HTML twin (for high-fidelity PDF rendering via Chromium) ──────────────────
def _article_html(item, lang, lead=False):
    paras = ''.join(f'<p class="body">{_html.escape(b)}</p>' for b in item['body'][lang])
    return f'''<article class="{'lead' if lead else 'item'}">
      <div class="num">{item['n']}</div>
      <div class="kicker">{_html.escape(item['kicker'][lang])}</div>
      <h2 class="headline">{_html.escape(item['headline'][lang])}</h2>
      {paras}
      <p class="source">{_html.escape(item['source'][lang])}</p>
    </article>'''


def build_html(lang, basename, chrome, lead, articles, outdir):
    ch = chrome[lang]
    doc = f'''<!DOCTYPE html><html lang="{'az' if lang=='AZ' else 'en'}"><head><meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 12mm 10mm; }}
  * {{ box-sizing: border-box; }}
  body {{ margin: 0; color: #000;
    font-family: "Eloquia Text ExtLt", "Segoe UI", "Helvetica Neue", Arial, sans-serif; }}
  .masthead {{ display: flex; align-items: baseline; gap: 10px;
    border-bottom: 1.5px solid #1F3864; padding-bottom: 4px; }}
  .masthead .omni {{ font-size: 54pt; letter-spacing: -1px; }}
  .masthead .lawg {{ font-size: 17pt; }}
  .masthead .lawg b {{ color: #1F3864; }}
  .banner {{ display: flex; justify-content: space-between; align-items: baseline;
    border-bottom: 0.75px solid #1F3864; padding: 4px 0 5px; margin-bottom: 10px; }}
  .banner .left {{ font-size: 9pt; font-weight: 700; color: #1F3864; letter-spacing: 0.3px; }}
  .banner .right {{ font-size: 9pt; color: #3B3838; }}
  .num {{ text-align: center; font-weight: 700; color: #3B3838; line-height: 1; }}
  .item .num {{ font-size: 26pt; margin-top: 10px; }}
  .lead .num {{ font-size: 36pt; margin-top: 4px; }}
  .kicker {{ text-align: center; font-family: "Times New Roman", Georgia, serif;
    font-size: 7pt; font-weight: 700; color: #806000; letter-spacing: 0.4px; margin-top: 2px; }}
  .headline {{ text-align: center; font-family: "Times New Roman", Georgia, serif;
    font-weight: 700; color: #1F3864; margin: 2px 0 6px; }}
  .item .headline {{ font-size: 12pt; line-height: 1.2; }}
  .lead .headline {{ font-size: 14pt; }}
  .body {{ text-align: justify; font-size: 9pt; line-height: 1.45; margin: 0 0 6px; }}
  .source {{ text-align: justify; font-family: "Times New Roman", Georgia, serif;
    font-size: 8pt; font-style: italic; color: #8A6D1E; margin: 0 0 4px; }}
  .lead {{ margin-bottom: 6px; }}
  .columns {{ column-count: 3; column-gap: 14px; column-rule: 0.5px solid #d9d9d9; }}
  .columns .item {{ break-inside: avoid-column; -webkit-column-break-inside: avoid; }}
  .footer {{ border-top: 0.75px solid #1F3864; margin-top: 10px; padding-top: 5px; }}
  .footer .d {{ text-align: justify; font-size: 7.5pt; color: #3B3838; }}
  .footer .d b {{ color: #1F3864; }}
  .footer .c {{ font-size: 7.5pt; font-weight: 700; color: #3B3838; margin-top: 3px; }}
</style></head><body>
  <div class="masthead"><span class="omni">omni</span><span class="lawg">law <b>gazette</b></span></div>
  <div class="banner"><span class="left">{_html.escape(ch['banner'])}</span>
    <span class="right">{_html.escape(ch['issue'])}</span></div>
  {_article_html(lead, lang, lead=True)}
  <div class="columns">
    {''.join(_article_html(a, lang) for a in articles)}
  </div>
  <div class="footer">
    <p class="d"><b>{_html.escape(ch['footer_lead'])}</b>{_html.escape(ch['footer'])}</p>
    <p class="c">{_html.escape(ch['footer_c'])}</p>
  </div>
</body></html>'''
    out = os.path.join(outdir, f'{basename}-{lang}.html')
    with open(out, 'w', encoding='utf-8') as f:
        f.write(doc)
    return out


def build_issue(basename, chrome, lead, articles, outdir, langs=LANGS):
    os.makedirs(outdir, exist_ok=True)
    made = []
    for lang in langs:
        made.append(build_docx(lang, basename, chrome, lead, articles, outdir))
        made.append(build_html(lang, basename, chrome, lead, articles, outdir))
    return made
