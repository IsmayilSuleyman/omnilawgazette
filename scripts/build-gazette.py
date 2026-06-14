#!/usr/bin/env python3
"""Build the Omni Law Gazette as two separate, mirror-image editions.

Per the firm's editorial rule (see AGENTS.md), every issue is produced as two
distinct documents — one entirely in Azerbaijani, one entirely in English —
identical in layout, item order and content; only the language differs.

AUTHORING ORDER (important): Azerbaijani is the source-of-truth edition. The
official portal material (e-qanun.az, etc.) is already in Azerbaijani, so each
article's 'AZ' text is authored FIRST, faithful to the source legal wording;
the 'EN' text is a translation of the finished Azerbaijani. In the per-article
data below, treat 'AZ' as canonical and 'EN' as its translation.

For each language this script emits:
  * a native, editable .docx (Eloquia Text ExtLt typography)
  * a matching .html twin (used to render a high-fidelity .pdf via Chromium)

Layout (from the firm's example edition): A4 / 0.5in margins; single-column
masthead + full-width lead (item 01); three-column body of numbered narrative
articles; single-column footer.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import html as _html

# ── Palette / fonts ──────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1F, 0x38, 0x64)
GREY  = RGBColor(0x3B, 0x38, 0x38)
GOLD  = RGBColor(0x80, 0x60, 0x00)
GOLD2 = RGBColor(0x8A, 0x6D, 0x1E)
BLACK = RGBColor(0x00, 0x00, 0x00)
BODY_FONT = 'Eloquia Text ExtLt'
SER_FONT  = 'Times New Roman'

LANGS = ['EN', 'AZ']

# ── Per-language chrome (masthead/banner/footer) ─────────────────────────────
CHROME = {
    'EN': {
        'banner': 'WEEKLY LEGISLATIVE DIGEST · REPUBLIC OF AZERBAIJAN',
        'issue':  'ISSUE No. 1   ·   5–12 JUNE 2026',
        'footer_lead': 'Omni Law Gazette',
        'footer': '  ·  Internal weekly digest of the legislative activity of the Republic of '
                  'Azerbaijan, compiled by Omni Law Firm for informational purposes only. It does not '
                  'constitute legal advice; consult the consolidated texts on e-qanun.az before relying '
                  'on any item. Sources: president.az · nk.gov.az · meclis.gov.az · e-qanun.az · constcourt.gov.az',
        'footer_c': 'Issue No. 1  ·  Reporting period 5–12 June 2026  ·  © 2026 Omni Law Firm',
    },
    'AZ': {
        'banner': 'HƏFTƏLİK QANUNVERİCİLİK İCMALI · AZƏRBAYCAN RESPUBLİKASI',
        'issue':  'BURAXILIŞ № 1   ·   5–12 İYUN 2026',
        'footer_lead': 'Omni Law Gazette',
        'footer': '  ·  Azərbaycan Respublikasının qanunvericilik fəaliyyətinin həftəlik daxili '
                  'icmalı; Omni Law Firm tərəfindən yalnız məlumat məqsədilə hazırlanıb. Hüquqi məsləhət '
                  'təşkil etmir; istənilən maddəyə istinad etməzdən əvvəl e-qanun.az saytındakı konsolidə '
                  'edilmiş mətnlərə baxın. Mənbələr: president.az · nk.gov.az · meclis.gov.az · e-qanun.az · constcourt.gov.az',
        'footer_c': 'Buraxılış № 1  ·  Baxış dövrü 5–12 iyun 2026  ·  © 2026 Omni Law Firm',
    },
}

# ── Articles (14 instruments, editorially ranked) ────────────────────────────
LEAD = {
    'n': '01',
    'kicker': {'EN': 'PRESIDENT · DECREE · 9 JUNE 2026',
               'AZ': 'PREZİDENT · FƏRMAN · 9 İYUN 2026'},
    'headline': {'EN': 'President expands credit guarantees and interest subsidies for entrepreneurs.',
                 'AZ': 'Prezident sahibkarlar üçün kredit zəmanətlərini və faiz subsidiyalarını genişləndirir.'},
    'body': {
        'EN': [
            "The week’s defining act is a presidential decree, signed 9 June, that overhauls the "
            "state’s system of credit support for the private sector. The decree refines the mechanisms "
            "through which the government guarantees bank loans extended to entrepreneurs and subsidises the "
            "interest payable on them, and it amends a series of earlier presidential decrees so that the whole "
            "framework operates under a single, updated set of rules.",
            "In practice, the reform widens access to state-backed financing and lowers the effective cost of "
            "borrowing for private business. By reducing the collateral burden and the interest expense that "
            "small and medium enterprises face, it is designed to channel commercial credit into productive, "
            "non-oil sectors of the economy — manufacturing, agriculture, services and export-oriented production.",
            "It is the most consequential of the fourteen instruments entered into the official State Register "
            "during the review period, and it sits at the centre of a wider economic-policy package. The same day, "
            "the President signed a complementary decree on non-oil-gas exports (item 02), and on 8 June overhauled "
            "the Agrarian Credit and Development Agency (item 03) — three measures that together point firmly "
            "toward diversification away from hydrocarbon revenue.",
        ],
        'AZ': [
            "Həftənin əsas hadisəsi 9 iyunda imzalanmış, özəl sektora dövlət kredit dəstəyi sistemini "
            "əsaslı şəkildə yeniləyən prezident fərmanıdır. Fərman sahibkarlara verilən bank kreditlərinə dövlət "
            "zəmanəti verilməsi və bu kreditlər üzrə faizlərin subsidiyalaşdırılması mexanizmlərini təkmilləşdirir "
            "və bütün çərçivənin vahid, yenilənmiş qaydalar əsasında işləməsi üçün prezidentin bir sıra əvvəlki "
            "fərmanlarına dəyişikliklər edir.",
            "Praktikada islahat dövlət dəstəkli maliyyəyə çıxışı genişləndirir və özəl biznes üçün borclanmanın "
            "real dəyərini azaldır. Kiçik və orta sahibkarlığın üzləşdiyi girov yükünü və faiz xərclərini "
            "azaltmaqla, o, kommersiya kreditini iqtisadiyyatın məhsuldar, qeyri-neft sektorlarına — istehsala, "
            "kənd təsərrüfatına, xidmətlərə və ixracyönümlü istehsala — yönəltməyi hədəfləyir.",
            "Bu, baxış dövründə rəsmi Dövlət Reyestrinə daxil edilmiş on dörd sənəddən ən əhəmiyyətlisidir və "
            "daha geniş iqtisadi siyasət paketinin mərkəzində dayanır. Həmin gün Prezident qeyri-neft-qaz ixracına "
            "dair əlavə fərman (maddə 02), 8 iyunda isə Aqrar Kredit və İnkişaf Agentliyini yenidən təşkil edən "
            "fərman (maddə 03) imzalayıb — birlikdə bu üç tədbir karbohidrogen gəlirlərindən diversifikasiyaya "
            "doğru qətiyyətli istiqaməti göstərir.",
        ],
    },
    'source': {'EN': 'Source · e-qanun.az/framework/62002',
               'AZ': 'Mənbə · e-qanun.az/framework/62002'},
}

ARTICLES = [
    {
        'n': '02',
        'kicker': {'EN': 'PRESIDENT · DECREE · 9 JUNE 2026', 'AZ': 'PREZİDENT · FƏRMAN · 9 İYUN 2026'},
        'headline': {'EN': 'A parallel push to reward non-oil exporters.',
                     'AZ': 'Qeyri-neft ixracatçılarını həvəsləndirmək üçün paralel addım.'},
        'body': {
            'EN': [
                "Signed the same day as the credit-guarantee reform, a second presidential decree introduces "
                "additional measures to stimulate the export of non-oil-and-gas goods and to defray the transport "
                "and logistics costs that weigh on exporters. The instrument extends targeted support to producers "
                "competing in foreign markets, where freight expense often erodes the margin on Azerbaijani goods.",
                "Read together with item 01, the decree confirms a deliberate tilt of fiscal policy toward the "
                "tradable, non-hydrocarbon economy. For manufacturers, agro-processors and logistics operators, the "
                "practical question will be the eligibility thresholds and the per-shipment ceilings, which fall to "
                "be set in the implementing rules.",
            ],
            'AZ': [
                "Kredit zəmanəti islahatı ilə eyni gün imzalanan ikinci prezident fərmanı qeyri-neft-qaz mallarının "
                "ixracının stimullaşdırılması və ixracatçıların üzərinə düşən nəqliyyat və logistika xərclərinin "
                "qismən ödənilməsi üçün əlavə tədbirlər tətbiq edir. Sənəd xarici bazarlarda rəqabət aparan "
                "istehsalçılara — daşıma xərclərinin çox vaxt Azərbaycan mallarının mənfəətini azaltdığı bir "
                "şəraitdə — hədəflənmiş dəstək göstərir.",
                "Maddə 01 ilə birlikdə oxunduqda, fərman fiskal siyasətin ticari, qeyri-karbohidrogen "
                "iqtisadiyyatına məqsədyönlü yönəlməsini təsdiqləyir. İstehsalçılar, aqro-emalçılar və logistika "
                "operatorları üçün praktiki sual icra qaydalarında müəyyən ediləcək uyğunluq həddləri və hər "
                "göndəriş üzrə yuxarı limitlər olacaq.",
            ],
        },
        'source': {'EN': 'Source · president.az/az/articles/view/72713',
                   'AZ': 'Mənbə · president.az/az/articles/view/72713'},
    },
    {
        'n': '03',
        'kicker': {'EN': 'PRESIDENT · DECREE · 8 JUNE 2026', 'AZ': 'PREZİDENT · FƏRMAN · 8 İYUN 2026'},
        'headline': {'EN': 'The Agrarian Credit and Development Agency is recast.',
                     'AZ': 'Aqrar Kredit və İnkişaf Agentliyi yenidən qurulur.'},
        'body': {
            'EN': [
                "On 8 June the President amended the founding statute of the Agrarian Credit and Development Agency "
                "under the Ministry of Agriculture — the 2015 decree (No. 571) that defines the body — "
                "together with related instruments governing state support to farming. The changes refit the agency "
                "that channels concessional finance into the rural economy.",
                "Coming a day before the broader entrepreneur-credit reform, the measure signals that agriculture is "
                "to be a primary beneficiary of the diversification drive. Agribusiness clients should expect revised "
                "lending criteria and, potentially, a wider mandate for the agency.",
            ],
            'AZ': [
                "8 iyunda Prezident Kənd Təsərrüfatı Nazirliyi yanında Aqrar Kredit və İnkişaf Agentliyinin təsis "
                "sənədini — qurumu müəyyən edən 2015-ci il tarixli fərmanı (№ 571) — və kənd təsərrüfatına dövlət "
                "dəstəyini tənzimləyən əlaqəli sənədləri dəyişdirib. Dəyişikliklər kənd iqtisadiyyatına güzəştli "
                "maliyyə yönəldən agentliyi yenidən nizamlayır.",
                "Daha geniş sahibkarlıq krediti islahatından bir gün əvvəl gələn tədbir kənd təsərrüfatının "
                "diversifikasiya kursunun əsas faydalananlarından biri olacağını göstərir. Aqrobiznes müştəriləri "
                "yenilənmiş kreditləşmə meyarlarını və ehtimal ki, agentliyin səlahiyyətlərinin genişlənməsini "
                "gözləməlidirlər.",
            ],
        },
        'source': {'EN': 'Source · president.az/az/articles/view/72706',
                   'AZ': 'Mənbə · president.az/az/articles/view/72706'},
    },
    {
        'n': '04',
        'kicker': {'EN': 'PRESIDENT · DECREE · 12 JUNE 2026', 'AZ': 'PREZİDENT · FƏRMAN · 12 İYUN 2026'},
        'headline': {'EN': 'A single rulebook for the country’s airports.',
                     'AZ': 'Ölkənin hava limanları üçün vahid qaydalar toplusu.'},
        'body': {
            'EN': [
                "The President approved, on 12 June, a comprehensive set of Rules on the Management of Airports "
                "(Aerodromes), issued under the constitutional power to regulate (Article 109.32) and framed as part "
                "of measures to improve management in civil aviation. The Rules establish the governance model, "
                "operational standards and responsibilities that will bind airport operators.",
                "For parties to concession, ground-handling, leasing or infrastructure agreements at the country’s "
                "airports, the decree introduces a new compliance baseline. Existing contracts should be reviewed "
                "against the Rules, and any transitional provisions noted.",
            ],
            'AZ': [
                "12 iyunda Prezident Konstitusiyanın 109-cu maddəsinin 32-ci bəndinə əsasən və mülki aviasiyada "
                "idarəetmənin təkmilləşdirilməsi tədbirlərinin bir hissəsi kimi “Hava limanlarının (aeroportların) "
                "idarə edilməsi Qaydası”nı təsdiq edib. Qaydalar hava limanı operatorları üçün məcburi olan "
                "idarəetmə modelini, əməliyyat standartlarını və məsuliyyətləri müəyyən edir.",
                "Ölkənin hava limanlarında konsessiya, yerüstü xidmət, icarə və ya infrastruktur müqavilələrinin "
                "tərəfləri üçün fərman yeni uyğunluq bazası tətbiq edir. Mövcud müqavilələr Qaydalar baxımından "
                "nəzərdən keçirilməli və hər hansı keçid müddəaları qeydə alınmalıdır.",
            ],
        },
        'source': {'EN': 'Source · president.az/az/articles/view/72735',
                   'AZ': 'Mənbə · president.az/az/articles/view/72735'},
    },
    {
        'n': '05',
        'kicker': {'EN': 'PRESIDENT · DECREE · 10 JUNE 2026', 'AZ': 'PREZİDENT · FƏRMAN · 10 İYUN 2026'},
        'headline': {'EN': 'The capital’s administration is restructured.',
                     'AZ': 'Paytaxtın idarəetməsi yenidən qurulur.'},
        'body': {
            'EN': [
                "A decree of 10 June reorganises the structure and management of the Baku City Executive Authority, "
                "the body that administers the capital. The reform adjusts the internal architecture and lines of "
                "management of the country’s largest municipal administration.",
                "Governance reforms of this kind tend to ripple outward — into procurement, permitting and the "
                "handling of municipal matters generally. Clients with dealings before the Baku authority should watch "
                "for consequential changes to competences and points of contact.",
            ],
            'AZ': [
                "10 iyun tarixli fərman paytaxtı idarə edən qurum olan Bakı Şəhər İcra Hakimiyyətinin strukturunu "
                "və idarə edilməsini yenidən təşkil edir. İslahat ölkənin ən böyük bələdiyyə idarəçiliyinin daxili "
                "quruluşunu və idarəetmə xəttini tənzimləyir.",
                "Bu cür idarəetmə islahatları adətən daha geniş təsir göstərir — satınalmalara, icazələrin "
                "verilməsinə və bələdiyyə məsələlərinin həllinə. Bakı İcra Hakimiyyəti qarşısında işləri olan "
                "müştərilər səlahiyyətlərdə və əlaqə nöqtələrində mümkün dəyişiklikləri izləməlidirlər.",
            ],
        },
        'source': {'EN': 'Source · president.az/az/articles/view/72724',
                   'AZ': 'Mənbə · president.az/az/articles/view/72724'},
    },
    {
        'n': '06',
        'kicker': {'EN': 'PRESIDENT · DECREE · 10 JUNE 2026', 'AZ': 'PREZİDENT · FƏRMAN · 10 İYUN 2026'},
        'headline': {'EN': 'The new citizens’-appeals law takes effect.',
                     'AZ': 'Vətəndaşların müraciətlərinə dair yeni qanun qüvvəyə minir.'},
        'body': {
            'EN': [
                "The President brought into force the 26 May amendments to the Law on Citizens’ Appeals "
                "(No. 415-VIIQD) and, in the same decree, amended the 2016 Rules (approved by Decree No. 950) on the "
                "handling of citizens’ appeals across state and municipal bodies, state-controlled legal entities "
                "and budget organisations.",
                "The measure tightens the procedural framework that governs how public bodies receive and answer "
                "citizens. Administrative-law practitioners should note the revised record-keeping and response "
                "obligations, which take effect immediately.",
            ],
            'AZ': [
                "Prezident “Vətəndaşların müraciətləri haqqında” Qanuna 26 may tarixli dəyişiklikləri "
                "(№ 415-VIIQD) qüvvəyə mindirib və eyni fərmanla dövlət və bələdiyyə orqanlarında, dövlət "
                "nəzarətindəki hüquqi şəxslərdə və büdcə təşkilatlarında vətəndaşların müraciətləri ilə bağlı "
                "kargüzarlığa dair 2016-cı il Qaydalarına (№ 950 nömrəli fərmanla təsdiq edilmiş) dəyişikliklər edib.",
                "Tədbir dövlət orqanlarının vətəndaşları qəbul etmə və onlara cavab vermə proseduru çərçivəsini "
                "sıxlaşdırır. İnzibati hüquq üzrə praktiklər dərhal qüvvəyə minən yenilənmiş kargüzarlıq və cavab "
                "öhdəliklərini nəzərə almalıdırlar.",
            ],
        },
        'source': {'EN': 'Source · president.az/az/articles/view/72722',
                   'AZ': 'Mənbə · president.az/az/articles/view/72722'},
    },
    {
        'n': '07',
        'kicker': {'EN': 'CABINET · DECISION No. 171 · 9 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · QƏRAR № 171 · 9 İYUN 2026'},
        'headline': {'EN': 'Rules adopted for the national rail network.',
                     'AZ': 'Milli dəmir yolu şəbəkəsi üçün qaydalar qəbul edilir.'},
        'body': {
            'EN': [
                "The Cabinet of Ministers approved Rules on the maintenance and use of general-use railway transport "
                "infrastructure — a foundational regulation for the rail sector. The Rules govern how the publicly "
                "accessible network is kept up and accessed by operators.",
                "For freight forwarders, infrastructure contractors and any party to a rail-access arrangement, the "
                "decision sets the terms of engagement with the network. It is the most structurally significant of "
                "the Cabinet’s eleven decisions this week.",
            ],
            'AZ': [
                "Nazirlər Kabineti ümumi istifadədə olan dəmiryol nəqliyyatı infrastrukturunun saxlanması və ondan "
                "istifadə Qaydalarını təsdiq edib — bu, dəmir yolu sektoru üçün təməl tənzimləmədir. Qaydalar "
                "ictimai əlçatan şəbəkənin necə saxlanılacağını və operatorlar tərəfindən necə istifadə ediləcəyini "
                "müəyyən edir.",
                "Yük daşıyıcıları, infrastruktur podratçıları və dəmir yoluna çıxış razılaşmasının istənilən tərəfi "
                "üçün qərar şəbəkə ilə qarşılıqlı əlaqənin şərtlərini müəyyən edir. Bu, Kabinetin həmin həftədəki on "
                "bir qərarından struktur baxımından ən əhəmiyyətlisidir.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Decision No. 171',
                   'AZ': 'Mənbə · nk.gov.az · Qərar № 171'},
    },
    {
        'n': '08',
        'kicker': {'EN': 'CABINET · DECISION No. 173 · 9 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · QƏRAR № 173 · 9 İYUN 2026'},
        'headline': {'EN': 'The gas-supply rulebook is rewritten.',
                     'AZ': 'Qaz təchizatı qaydaları yenidən yazılır.'},
        'body': {
            'EN': [
                "Implementing the 2025 Law on Gas Supply (No. 233-VIIQ), the Cabinet amended a series of its own "
                "decisions and repealed three legacy instruments — rules dating to 1999 and 2023 on gas "
                "distribution and on the construction and operation of gas installations. The effect is to consolidate "
                "the sector under the current statutory regime.",
                "Distributors and large industrial consumers should map the repealed rules against the new framework "
                "to confirm that internal compliance procedures remain valid.",
            ],
            'AZ': [
                "2025-ci il “Qaz təchizatı haqqında” Qanunun (№ 233-VIIQ) icrası ilə əlaqədar Nazirlər Kabineti öz "
                "qərarlarının bir sırasını dəyişdirib və üç köhnə sənədi — qaz paylanmasına, habelə qaz qurğularının "
                "tikintisinə və istismarına dair 1999 və 2023-cü illərə aid qaydaları — ləğv edib. Nəticədə sektor "
                "cari qanunvericilik rejimi altında birləşdirilir.",
                "Paylayıcılar və iri sənaye istehlakçıları daxili uyğunluq prosedurlarının qüvvədə qaldığını "
                "təsdiqləmək üçün ləğv edilmiş qaydaları yeni çərçivə ilə tutuşdurmalıdırlar.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Decision No. 173',
                   'AZ': 'Mənbə · nk.gov.az · Qərar № 173'},
    },
    {
        'n': '09',
        'kicker': {'EN': 'CABINET · DECISION No. 174 · 10 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · QƏRAR № 174 · 10 İYUN 2026'},
        'headline': {'EN': 'State guarantees underpin a new upstream PSA.',
                     'AZ': 'Dövlət zəmanətləri yeni hasilat pay bölgüsü sazişini dəstəkləyir.'},
        'body': {
            'EN': [
                "The Cabinet approved the Government’s guarantees and obligations under the production-sharing "
                "agreement between SOCAR and Gran Tierra Energy (Azerbaijan) GmbH, covering exploration, development "
                "and production across parts of the Quba, Qusar, Shabran, Siyazan and Khizi districts. The decision "
                "gives the sovereign backing that international upstream investors require.",
                "The award signals continued foreign-investor appetite for Azerbaijani acreage and provides a "
                "template for the contractual architecture of future production-sharing agreements.",
            ],
            'AZ': [
                "Nazirlər Kabineti SOCAR ilə “Gran Tierra Energy (Azerbaijan) GmbH” arasında Quba, Qusar, Şabran, "
                "Siyəzən və Xızı rayonlarının bəzi hissələrini əhatə edən kəşfiyyat, işlənmə və hasilatın pay bölgüsü "
                "Sazişi üzrə Hökumətin zəmanət və öhdəliklərini təsdiq edib. Qərar beynəlxalq hasilat investorlarının "
                "tələb etdiyi suveren təminatı verir.",
                "Saziş Azərbaycan yataqlarına xarici investor marağının davam etdiyini göstərir və gələcək hasilat "
                "pay bölgüsü sazişlərinin müqavilə strukturu üçün nümunə təqdim edir.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Decision No. 174',
                   'AZ': 'Mənbə · nk.gov.az · Qərar № 174'},
    },
    {
        'n': '10',
        'kicker': {'EN': 'CABINET · DECISIONS Nos. 176–178 · 12 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · QƏRARLAR № 176–178 · 12 İYUN 2026'},
        'headline': {'EN': 'Three new sites earmarked for renewable energy.',
                     'AZ': 'Bərpa olunan enerji üçün üç yeni ərazi ayrılır.'},
        'body': {
            'EN': [
                "In a single sitting on 12 June the Cabinet designated two state-owned parcels in the Hacıqabul "
                "district as renewable-energy sites (Decisions Nos. 177 and 178) and amended its 2024 decision on a "
                "comparable site in Jabrayil (No. 176). The instruments enlarge the land bank available for solar and "
                "wind development.",
                "The Jabrayil amendment, revisiting a 2024 designation, shows the liberated-territories energy "
                "programme being refined in practice. Developers and lenders active in these districts should track "
                "the evolving site parameters.",
            ],
            'AZ': [
                "12 iyunda bir iclasda Nazirlər Kabineti Hacıqabul rayonunda iki dövlət torpaq sahəsini bərpa olunan "
                "enerji əraziləri kimi müəyyən edib (№ 177 və 178 saylı qərarlar) və Cəbrayıl rayonunda oxşar sahəyə "
                "dair 2024-cü il qərarına dəyişiklik edib (№ 176). Sənədlər günəş və külək enerjisinin inkişafı üçün "
                "mövcud torpaq ehtiyatını artırır.",
                "2024-cü il təyinatına yenidən baxan Cəbrayıl dəyişikliyi azad edilmiş ərazilərdə enerji proqramının "
                "praktikada təkmilləşdirildiyini göstərir. Bu rayonlarda fəal olan tərtibatçılar və kreditorlar "
                "dəyişən sahə parametrlərini izləməlidirlər.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Decisions No. 176–178',
                   'AZ': 'Mənbə · nk.gov.az · Qərarlar № 176–178'},
    },
    {
        'n': '11',
        'kicker': {'EN': 'CABINET · DECISION No. 169 · 9 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · QƏRAR № 169 · 9 İYUN 2026'},
        'headline': {'EN': 'The security service gains a central analytics system.',
                     'AZ': 'Təhlükəsizlik xidməti mərkəzləşdirilmiş analitika sistemi əldə edir.'},
        'body': {
            'EN': [
                "The Cabinet amended several of its decisions to give effect to a 2025 presidential decree (No. 533) "
                "establishing the Centralised Information and Digital Analytics System of the State Security Service. "
                "The decision aligns subordinate regulation with the new architecture for state-security data.",
                "While operationally sensitive, the measure forms part of a broader digitisation of the security "
                "apparatus and may bear on data-sharing obligations imposed on regulated entities.",
            ],
            'AZ': [
                "Nazirlər Kabineti Dövlət Təhlükəsizliyi Xidmətinin Mərkəzləşdirilmiş İnformasiya və Rəqəmsal "
                "Analitika Sisteminin yaradılmasına dair 2025-ci il prezident fərmanının (№ 533) icrası üçün öz "
                "qərarlarının bir neçəsinə dəyişiklik edib. Qərar tabe tənzimləməni dövlət təhlükəsizliyi "
                "məlumatları üçün yeni arxitektura ilə uzlaşdırır.",
                "Əməliyyat baxımından həssas olsa da, tədbir təhlükəsizlik aparatının daha geniş rəqəmsallaşmasının "
                "bir hissəsidir və tənzimlənən qurumlara qoyulan məlumat mübadiləsi öhdəliklərinə təsir göstərə bilər.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Decision No. 169',
                   'AZ': 'Mənbə · nk.gov.az · Qərar № 169'},
    },
    {
        'n': '12',
        'kicker': {'EN': 'CABINET · DECISION No. 170 · 9 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · QƏRAR № 170 · 9 İYUN 2026'},
        'headline': {'EN': 'The mandatory-insurance benefit package is adjusted.',
                     'AZ': 'İcbari sığorta üzrə xidmətlər zərfi dəyişdirilir.'},
        'body': {
            'EN': [
                "The Cabinet amended the Mandatory Health Insurance Services Package — the schedule of services "
                "covered under the compulsory scheme, first approved in 2020. Changes to the package alter the scope "
                "of care that insured persons may claim and the corresponding obligations on providers.",
                "Employers and healthcare providers should confirm how the revised schedule affects coverage and "
                "reimbursement before the change takes effect.",
            ],
            'AZ': [
                "Nazirlər Kabineti İcbari Tibbi Sığorta üzrə Xidmətlər Zərfini — ilk dəfə 2020-ci ildə təsdiq "
                "edilmiş, icbari sxem çərçivəsində əhatə olunan xidmətlər siyahısını — dəyişdirib. Zərfdəki "
                "dəyişikliklər sığortalı şəxslərin tələb edə biləcəyi tibbi xidmətlərin əhatəsini və müvafiq olaraq "
                "xidmət təminatçılarının öhdəliklərini dəyişir.",
                "İşəgötürənlər və səhiyyə təminatçıları dəyişiklik qüvvəyə minməzdən əvvəl yenilənmiş siyahının "
                "əhatəyə və ödənişlərə necə təsir etdiyini dəqiqləşdirməlidirlər.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Decision No. 170',
                   'AZ': 'Mənbə · nk.gov.az · Qərar № 170'},
    },
    {
        'n': '13',
        'kicker': {'EN': 'CABINET · DECISION No. 168 · 9 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · QƏRAR № 168 · 9 İYUN 2026'},
        'headline': {'EN': 'A framework for energy-efficiency data.',
                     'AZ': 'Enerji səmərəliliyi məlumatları üçün çərçivə.'},
        'body': {
            'EN': [
                "The Cabinet approved the Regulation on the Energy Efficiency Information System, establishing the "
                "legal basis for collecting, monitoring and reporting energy-consumption data across the economy. The "
                "system is intended to underpin efficiency policy with reliable measurement.",
                "Energy-intensive enterprises should anticipate reporting duties once the system becomes operational.",
            ],
            'AZ': [
                "Nazirlər Kabineti “Enerji effektivliyi informasiya sisteminin Əsasnaməsi”ni təsdiq edərək "
                "iqtisadiyyat üzrə enerji istehlakı məlumatlarının toplanması, izlənməsi və hesabatı üçün hüquqi əsas "
                "yaradıb. Sistem səmərəlilik siyasətini etibarlı ölçmə ilə dəstəkləmək üçün nəzərdə tutulub.",
                "Enerji tutumlu müəssisələr sistem işə düşdükdən sonra hesabat öhdəliklərini gözləməlidirlər.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Decision No. 168',
                   'AZ': 'Mənbə · nk.gov.az · Qərar № 168'},
    },
    {
        'n': '14',
        'kicker': {'EN': 'CABINET · ORDERS · 9–10 JUNE 2026',
                   'AZ': 'NAZİRLƏR KABİNETİ · SƏRƏNCAMLAR · 9–10 İYUN 2026'},
        'headline': {'EN': 'Four orders put recent presidential decrees into effect.',
                     'AZ': 'Dörd sərəncam son prezident fərmanlarını icra edir.'},
        'body': {
            'EN': [
                "Alongside its decisions, the Cabinet issued four implementing orders. Two of 10 June give effect to "
                "presidential decrees on the ASAN citizen-service agency (No. 678) and on Finance Ministry "
                "representation on the boards of state-owned entities (No. 675). Two of 9 June implement the renaming "
                "of territorial units in the Quba district (under Law No. 404-VIIQ) and changes to the State Agency "
                "for the Protection of Strategic Facilities.",
                "Individually routine, together they show the machinery by which presidential policy is translated "
                "into administrative practice within days of signature.",
            ],
            'AZ': [
                "Qərarları ilə yanaşı, Nazirlər Kabineti dörd icra sərəncamı verib. 10 iyun tarixli ikisi ASAN "
                "xidmət agentliyinə (№ 678) və dövlət qurumlarının idarə heyətlərində Maliyyə Nazirliyinin "
                "təmsilçiliyinə (№ 675) dair prezident fərmanlarını icra edir. 9 iyun tarixli ikisi isə Quba "
                "rayonunda ərazi vahidlərinin adlarının dəyişdirilməsini (№ 404-VIIQ Qanunu əsasında) və Strateji "
                "Obyektlərin Mühafizəsi Dövlət Agentliyinə dair dəyişiklikləri tətbiq edir.",
                "Ayrı-ayrılıqda adi görünsə də, birlikdə onlar prezident siyasətinin imzalanmasından bir neçə gün "
                "sonra inzibati praktikaya necə çevrildiyi mexanizmini göstərir.",
            ],
        },
        'source': {'EN': 'Source · nk.gov.az · Orders No. 389s, 400s, 402s, 403s',
                   'AZ': 'Mənbə · nk.gov.az · Sərəncamlar № 389s, 400s, 402s, 403s'},
    },
]

OUTDIR = os.path.join(os.path.dirname(__file__), 'gazettes')
BASENAME = 'Omni-Law-Gazette-Issue-1-2026-06-12'


# ── DOCX helpers ─────────────────────────────────────────────────────────────
def set_cols(section, num):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), '708')


def run(doc, p, text, *, font=BODY_FONT, size=9, bold=False, italic=False, color=None):
    r = p.add_run(text)
    rPr = r._r.get_or_add_rPr(); rfonts = rPr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rPr.insert(0, rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), font)
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def para(doc, align=None, before=0, after=0, line=1.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    return p


def bottom_border(el, color='1F3864', sz='4'):
    pPr = el._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '4'); b.set(qn('w:color'), color)
    pbdr.append(b)


def build_docx(lang):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = BODY_FONT; st.font.size = Pt(9)
    st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.0
    ch = CHROME[lang]

    # masthead (brand stays "omni law gazette" in both editions)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, after=0)
    run(doc, p, 'omni', font=BODY_FONT, size=54, color=BLACK)
    run(doc, p, '  law ', font=BODY_FONT, size=17, color=BLACK)
    run(doc, p, 'gazette', font=BODY_FONT, size=17, color=NAVY)
    bottom_border(p, sz='6')

    # banner
    p = para(doc, after=8)
    run(doc, p, ch['banner'], font=BODY_FONT, size=9, bold=True, color=NAVY)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.27), WD_TAB_ALIGNMENT.RIGHT)
    run(doc, p, '\t', size=9)
    run(doc, p, ch['issue'], font=BODY_FONT, size=9, color=GREY)
    bottom_border(p, sz='4')
    para(doc, after=4)

    def render(item, lead=False):
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=(8 if lead else 12), after=0)
        run(doc, p, item['n'], font=BODY_FONT, size=(36 if lead else 28), bold=True, color=GREY)
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=1)
        run(doc, p, item['kicker'][lang], font=SER_FONT, size=7, bold=True, color=GOLD)
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
        run(doc, p, item['headline'][lang], font=SER_FONT, size=(13 if lead else 12), bold=True, color=NAVY)
        for par in item['body'][lang]:
            p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4, line=1.05)
            run(doc, p, par, font=BODY_FONT, size=9, color=BLACK)
        p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6)
        run(doc, p, item['source'][lang], font=SER_FONT, size=8, italic=True, color=GOLD2)

    render(LEAD, lead=True)
    doc.add_section(WD_SECTION.CONTINUOUS); set_cols(doc.sections[0], 1)
    for item in ARTICLES:
        render(item)
    doc.add_section(WD_SECTION.CONTINUOUS); set_cols(doc.sections[1], 3)

    p = para(doc, before=6, after=2); bottom_border(p, sz='4')
    p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=2)
    run(doc, p, ch['footer_lead'], font=BODY_FONT, size=7.5, bold=True, color=NAVY)
    run(doc, p, ch['footer'], font=BODY_FONT, size=7.5, color=GREY)
    p = para(doc, after=0)
    run(doc, p, ch['footer_c'], font=BODY_FONT, size=7.5, bold=True, color=GREY)
    set_cols(doc.sections[2], 1)

    for s in doc.sections:
        s.page_width = Inches(8.27); s.page_height = Inches(11.69)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.5)

    out = os.path.join(OUTDIR, f'{BASENAME}-{lang}.docx')
    doc.save(out)
    return out


# ── HTML twin ────────────────────────────────────────────────────────────────
def article_html(item, lang, lead=False):
    paras = ''.join(f'<p class="body">{_html.escape(b)}</p>' for b in item['body'][lang])
    return f'''<article class="{'lead' if lead else 'item'}">
      <div class="num">{item['n']}</div>
      <div class="kicker">{_html.escape(item['kicker'][lang])}</div>
      <h2 class="headline">{_html.escape(item['headline'][lang])}</h2>
      {paras}
      <p class="source">{_html.escape(item['source'][lang])}</p>
    </article>'''


def build_html(lang):
    ch = CHROME[lang]
    doc = f'''<!DOCTYPE html><html lang="{'az' if lang=='AZ' else 'en'}"><head><meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 12mm 10mm; }}
  * {{ box-sizing: border-box; }}
  body {{ margin: 0; color: #000;
    font-family: "Eloquia Text ExtLt", "Segoe UI", "Helvetica Neue", Arial, sans-serif; }}
  .masthead {{ display: flex; align-items: baseline; gap: 10px;
    border-bottom: 1.5px solid #1F3864; padding-bottom: 4px; }}
  .masthead .omni {{ font-size: 54pt; letter-spacing: -1px; }}
  .masthead .lawg {{ font-size: 17pt; }}
  .masthead .lawg b {{ color: #1F3864; }}
  .banner {{ display: flex; justify-content: space-between; align-items: baseline;
    border-bottom: 0.75px solid #1F3864; padding: 4px 0 5px; margin-bottom: 10px; }}
  .banner .left {{ font-size: 9pt; font-weight: 700; color: #1F3864; letter-spacing: 0.3px; }}
  .banner .right {{ font-size: 9pt; color: #3B3838; }}
  .num {{ text-align: center; font-weight: 700; color: #3B3838; line-height: 1; }}
  .item .num {{ font-size: 26pt; margin-top: 10px; }}
  .lead .num {{ font-size: 36pt; margin-top: 4px; }}
  .kicker {{ text-align: center; font-family: "Times New Roman", Georgia, serif;
    font-size: 7pt; font-weight: 700; color: #806000; letter-spacing: 0.4px; margin-top: 2px; }}
  .headline {{ text-align: center; font-family: "Times New Roman", Georgia, serif;
    font-weight: 700; color: #1F3864; margin: 2px 0 6px; }}
  .item .headline {{ font-size: 12pt; line-height: 1.2; }}
  .lead .headline {{ font-size: 14pt; }}
  .body {{ text-align: justify; font-size: 9pt; line-height: 1.45; margin: 0 0 6px; }}
  .source {{ text-align: justify; font-family: "Times New Roman", Georgia, serif;
    font-size: 8pt; font-style: italic; color: #8A6D1E; margin: 0 0 4px; }}
  .lead {{ margin-bottom: 6px; }}
  .columns {{ column-count: 3; column-gap: 14px; column-rule: 0.5px solid #d9d9d9; }}
  .columns .item {{ break-inside: avoid-column; -webkit-column-break-inside: avoid; }}
  .footer {{ border-top: 0.75px solid #1F3864; margin-top: 10px; padding-top: 5px; }}
  .footer .d {{ text-align: justify; font-size: 7.5pt; color: #3B3838; }}
  .footer .d b {{ color: #1F3864; }}
  .footer .c {{ font-size: 7.5pt; font-weight: 700; color: #3B3838; margin-top: 3px; }}
</style></head><body>
  <div class="masthead"><span class="omni">omni</span><span class="lawg">law <b>gazette</b></span></div>
  <div class="banner"><span class="left">{_html.escape(ch['banner'])}</span>
    <span class="right">{_html.escape(ch['issue'])}</span></div>
  {article_html(LEAD, lang, lead=True)}
  <div class="columns">
    {''.join(article_html(a, lang) for a in ARTICLES)}
  </div>
  <div class="footer">
    <p class="d"><b>{_html.escape(ch['footer_lead'])}</b>{_html.escape(ch['footer'])}</p>
    <p class="c">{_html.escape(ch['footer_c'])}</p>
  </div>
</body></html>'''
    out = os.path.join(OUTDIR, f'{BASENAME}-{lang}.html')
    with open(out, 'w', encoding='utf-8') as f:
        f.write(doc)
    return out


if __name__ == '__main__':
    os.makedirs(OUTDIR, exist_ok=True)
    for lang in LANGS:
        d = build_docx(lang)
        h = build_html(lang)
        print(f'[{lang}] DOCX -> {os.path.basename(d)} | HTML -> {os.path.basename(h)}')
