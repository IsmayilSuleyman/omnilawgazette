#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Omni Law Gazette — Word (.docx) builder, Issue No. 2

Renders a bilingual (Azerbaijani original / English translation) weekly
legislative digest for Omni Law Firm, covering the Republic of Azerbaijan's
Tier-1 official sources for the reporting period 15-22 June 2026.

Sources consulted at compilation:
  president.az/az/documents  — presidential decrees & orders (retrieved)
  nk.gov.az/az/senedler      — Cabinet of Ministers decisions & orders (retrieved)
  meclis.gov.az              — Milli Majlis bills (retrieved)
  constcourt.gov.az          — Constitutional Court (HTTP 503 at compilation)
  e-qanun.az                 — consolidated register (JS frameset; not list-scrapable)

Usage:  python3 scripts/build-gazette-docx.py
Output: scripts/gazettes/gazette-2026-06-22.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── palette ───────────────────────────────────────────────────────────────────
INK      = RGBColor(0x1a, 0x1a, 0x1a)
GREY     = RGBColor(0x55, 0x55, 0x55)
LGREY    = RGBColor(0x88, 0x88, 0x88)
DECREE   = RGBColor(0x2c, 0x3e, 0x50)
ORDER    = RGBColor(0x4a, 0x4a, 0x4a)
BILL     = RGBColor(0x6b, 0x6b, 0x6b)
LAW      = RGBColor(0x1a, 0x3a, 0x1a)
WHITE    = RGBColor(0xff, 0xff, 0xff)
AMBER    = RGBColor(0x6b, 0x59, 0x00)
SERIF    = "Georgia"
SANS     = "Arial"

# ── low-level xml helpers ─────────────────────────────────────────────────────
# OOXML enforces a strict child order inside <w:pPr>, <w:rPr> and <w:tcPr>.
# Word is lenient but LibreOffice is not, so we insert each element in the
# correct position using insert_element_before(tag, *successor_tags).
PPR_ORDER = ['w:pStyle','w:keepNext','w:keepLines','w:pageBreakBefore','w:framePr',
 'w:widowControl','w:numPr','w:suppressLineNumbers','w:pBdr','w:shd','w:tabs',
 'w:suppressAutoHyphens','w:kinsoku','w:wordWrap','w:overflowPunct','w:topLinePunct',
 'w:autoSpaceDE','w:autoSpaceDN','w:bidi','w:adjustRightInd','w:snapToGrid','w:spacing',
 'w:ind','w:contextualSpacing','w:mirrorIndents','w:suppressOverlap','w:jc',
 'w:textDirection','w:textAlignment','w:textboxTightWrap','w:outlineLvl','w:divId',
 'w:cnfStyle','w:rPr','w:sectPr','w:pPrChange']
RPR_ORDER = ['w:rStyle','w:rFonts','w:b','w:bCs','w:i','w:iCs','w:caps','w:smallCaps',
 'w:strike','w:dstrike','w:outline','w:shadow','w:emboss','w:imprint','w:noProof',
 'w:snapToGrid','w:vanish','w:webHidden','w:color','w:spacing','w:w','w:kern',
 'w:position','w:sz','w:szCs','w:highlight','w:u','w:effect','w:bdr','w:shd','w:fitText',
 'w:vertAlign','w:rtl','w:cs','w:em','w:lang','w:eastAsianLayout','w:specVanish','w:oMath']
TCPR_ORDER = ['w:cnfStyle','w:tcW','w:gridSpan','w:hMerge','w:vMerge','w:tcBorders',
 'w:shd','w:noWrap','w:tcMar','w:textDirection','w:tcFitText','w:vAlign','w:hideMark',
 'w:headers','w:cellIns','w:cellDel','w:cellMerge','w:tcPrChange']

def _insert_ordered(parent, tag, order):
    """Create <tag>, drop any existing one, insert it in schema position."""
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    el = OxmlElement(tag)
    successors = order[order.index(tag) + 1:]
    parent.insert_element_before(el, *successors)
    return el

def _shd(parent, hex_fill, order):
    shd = _insert_ordered(parent, 'w:shd', order)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)

def shade_paragraph(p, hex_fill):
    _shd(p._p.get_or_add_pPr(), hex_fill, PPR_ORDER)

def shade_run(r, hex_fill):
    _shd(r._r.get_or_add_rPr(), hex_fill, RPR_ORDER)

def shade_cell(cell, hex_fill):
    _shd(cell._tc.get_or_add_tcPr(), hex_fill, TCPR_ORDER)

def set_borders(p, sides=('bottom',), sz=6, color='1a1a1a', space=1):
    pPr = p._p.get_or_add_pPr()
    pbdr = _insert_ordered(pPr, 'w:pBdr', PPR_ORDER)
    for side in sides:
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), str(space))
        e.set(qn('w:color'), color)
        pbdr.append(e)

def no_space(p, before=0, after=0, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line

def run(p, text, *, size=10.5, bold=False, italic=False, color=INK,
        font=SERIF, caps=False, spacing=None):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    r.font.name = font
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        sp = _insert_ordered(r._r.get_or_add_rPr(), 'w:spacing', RPR_ORDER)
        sp.set(qn('w:val'), str(spacing))
    return r

# ── document scaffold ─────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.6); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
normal = doc.styles['Normal']
normal.font.name = SERIF
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK

# ── masthead ──────────────────────────────────────────────────────────────────
top = doc.add_paragraph(); no_space(top, 0, 2)
set_borders(top, sides=('top',), sz=24)
run(top, "INTERNAL CIRCULATION — OMNI LAW FIRM", size=7.5, color=GREY, font=SANS, caps=True)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(title, 4, 0)
run(title, "OMNI ", size=40, bold=True, font=SERIF, caps=True)
run(title, "Law", size=40, italic=True, font=SERIF)
run(title, " GAZETTE", size=40, bold=True, font=SERIF, caps=True)

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(sub, 2, 4)
run(sub, "Legislative Intelligence for Legal Professionals  ·  Hüquqşünaslar üçün Qanunvericilik İcmalı",
    size=9, color=GREY, font=SANS, caps=True, spacing=24)

meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(meta, 0, 4)
set_borders(meta, sides=('top', 'bottom'), sz=12)
run(meta, "ISSUE No. 2     ·     REPORTING PERIOD: 15–22 JUNE 2026     ·     PUBLISHED: MONDAY, 22 JUNE 2026",
    size=8, color=INK, font=SANS, caps=True)

banner = doc.add_paragraph(); banner.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(banner, 2, 10)
shade_paragraph(banner, '1a1a1a')
run(banner, "Week 25, 2026   ·   Republic of Azerbaijan   ·   Legislative Monitor",
    size=8.5, color=WHITE, font=SANS, caps=True, spacing=18)

# ── editor's note / at a glance ───────────────────────────────────────────────
def section_header(institution, title_en, title_az, right=""):
    h = doc.add_paragraph(); no_space(h, 14, 2)
    set_borders(h, sides=('top',), sz=18)
    run(h, institution.upper() + "  ", size=7, color=LGREY, font=SANS, spacing=20)
    t = doc.add_paragraph(); no_space(t, 0, 6)
    set_borders(t, sides=('bottom',), sz=6, color='cccccc')
    run(t, title_en, size=15, bold=True)
    run(t, "   " + title_az, size=10, italic=True, color=GREY)
    if right:
        run(t, "    — " + right, size=8, italic=True, color=LGREY, font=SANS)

def glance():
    h = doc.add_paragraph(); no_space(h, 4, 4)
    set_borders(h, sides=('left',), sz=24)
    shade_paragraph(h, 'fafafa')
    run(h, "BU HƏFTƏ BİR BAXIŞDA  /  THIS WEEK AT A GLANCE", size=8, color=GREY, font=SANS, caps=True, spacing=20)
    data = [("10", "Bills before\nParliament"),
            ("7",  "Presidential\nDecrees"),
            ("3",  "Presidential\nOrders"),
            ("6",  "Cabinet\nDecisions"),
            ("4",  "Cabinet\nOrders")]
    tbl = doc.add_table(rows=2, cols=len(data))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (num, label) in enumerate(data):
        c0 = tbl.cell(0, i).paragraphs[0]; c0.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(c0)
        run(c0, num, size=26, bold=True, color=INK)
        c1 = tbl.cell(1, i).paragraphs[0]; c1.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(c1)
        run(c1, label, size=8, color=GREY, font=SANS)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── document entry (bilingual) ────────────────────────────────────────────────
TAGCOLORS = {'DECREE': '2c3e50', 'ORDER': '4a4a4a', 'BILL': '6b6b6b',
             'DECISION': '2c3e50', 'LAW': '1a3a1a'}

def entry(tag, number, date, az_title, en_title, az_body, en_body):
    head = doc.add_paragraph(); no_space(head, 8, 1)
    r = head.add_run("  " + tag.upper() + "  ")
    r.font.size = Pt(6.5); r.font.name = SANS; r.bold = True
    r.font.color.rgb = WHITE
    shade_run(r, TAGCOLORS.get(tag.upper(), '1a1a1a'))
    if number:
        run(head, "  " + number, size=7.5, color=GREY, font=SANS)
    run(head, "      " + date, size=7.5, color=LGREY, font=SANS)

    t = doc.add_paragraph(); no_space(t, 0, 1, line=1.15)
    run(t, az_title, size=10, bold=True)
    e = doc.add_paragraph(); no_space(e, 0, 3, line=1.15)
    run(e, en_title, size=9, italic=True, color=RGBColor(0x33,0x33,0x55))

    ab = doc.add_paragraph(); no_space(ab, 0, 1, line=1.3)
    run(ab, "AZ  ", size=7, bold=True, color=LGREY, font=SANS)
    run(ab, az_body, size=9, color=RGBColor(0x22,0x22,0x22))
    eb = doc.add_paragraph(); no_space(eb, 0, 6, line=1.3)
    run(eb, "EN  ", size=7, bold=True, color=LGREY, font=SANS)
    run(eb, en_body, size=9, color=RGBColor(0x44,0x44,0x44))
    sep = doc.add_paragraph(); no_space(sep, 0, 0)
    set_borders(sep, sides=('bottom',), sz=4, color='d8d8d8')

def lead(az, en):
    p = doc.add_paragraph(); no_space(p, 2, 4, line=1.35)
    run(p, az, size=9, italic=True, color=GREY, font=SANS)
    p2 = doc.add_paragraph(); no_space(p2, 0, 8, line=1.35)
    run(p2, en, size=9, color=GREY, font=SANS)

def notice(kind, az, en):
    p = doc.add_paragraph(); no_space(p, 6, 6, line=1.3)
    fill = 'fff8e1' if kind == 'warn' else 'f4f4f4'
    bar  = 'c0a000' if kind == 'warn' else '1a1a1a'
    shade_paragraph(p, fill); set_borders(p, sides=('left',), sz=24, color=bar)
    run(p, az + "  ", size=8.5, color=AMBER if kind == 'warn' else GREY, font=SANS)
    run(p, en, size=8.5, color=AMBER if kind == 'warn' else GREY, font=SANS, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
glance()

intro = doc.add_paragraph(); no_space(intro, 2, 8, line=1.35)
run(intro, "REDAKTORDAN. ", size=9, bold=True, font=SANS, color=INK)
run(intro, "Həftə boyu fəaliyyət 16–19 iyun aralığında cəmləndi. 16 iyunda qeyri-neft ixracının dəstəklənməsinə dair Fərman və üç Nazirlər Kabineti qərarı, 17 iyunda iki təltif Sərəncamı, 19 iyunda isə altı Fərman, bir Sərəncam və Milli Məclisin on qanun layihəsi gündəmə gəldi. Diqqət mərkəzində diplomatik xidmət islahatı, vergi və investisiya təşviqi, qeyri-neft ixracı, enerji və rəqəmsal idarəetmə dayanır.",
    size=9, color=RGBColor(0x33,0x33,0x33))
intro2 = doc.add_paragraph(); no_space(intro2, 0, 6, line=1.35)
run(intro2, "FROM THE EDITORS. ", size=9, bold=True, font=SANS, color=INK)
run(intro2, "Activity clustered between 16 and 19 June. The 16th brought a Decree on non-oil export support and three Cabinet decisions; the 17th, two award Orders; and the 19th, six Decrees, one Order and ten bills tabled in the Milli Majlis. Dominant themes: diplomatic-service reform, tax and investment promotion, non-oil exports, energy, and digital governance. Items dated 16–17 June are drawn from the e-Qanun official register (with framework links below); the 19 June instruments are taken from president.az and nk.gov.az and had not yet been entered in the e-Qanun register at the time of compilation.",
    size=9, italic=True, color=GREY)

# ── SECTION 1: PRESIDENT ──────────────────────────────────────────────────────
section_header("Administration of the President · Azərbaycan Prezidentinin Administrasiyası",
               "Presidential Office", "Prezident İqamətgahı", "Decrees & Orders — 19 June 2026")
lead("Hesabat dövründə Prezident tərəfindən yeddi Fərman və üç Sərəncam dərc edilmişdir: 16 iyunda qeyri-neft ixracının dəstəklənməsi haqqında Fərman, 17 iyunda iki təltif Sərəncamı, 19 iyunda isə altı Fərman və bir Sərəncam. Sənədlər iqtisadiyyat, diplomatik xidmət, vergi-investisiya, təhsil maliyyələşməsi və enerji infrastrukturu sahələrini əhatə edir.",
     "Seven Decrees and three Orders were published by the President during the reporting period: a non-oil export-support Decree on 16 June, two award Orders on 17 June, and six Decrees and one Order on 19 June — spanning the economy, the diplomatic service, tax and investment, education financing, and energy infrastructure.")

entry("Decree", "e-Qanun id 62008", "16 Jun 2026",
      "Qeyri-neft-qaz malları üzrə ixracın dəstəklənməsi ilə bağlı əlavə tədbirlər haqqında",
      "On Additional Measures to Support the Export of Non-Oil-and-Gas Goods",
      "Fərman qeyri-neft-qaz sektorunda istehsal olunan malların ixracının stimullaşdırılması üçün əlavə dövlət dəstəyi tədbirlərini müəyyən edir. Tədbirlər ixrac təşviqi, qeyri-neft sahibkarlığının genişləndirilməsi və ixrac bazarlarının şaxələndirilməsi istiqamətində iqtisadi siyasətin davamıdır. Tam mətn e-Qanun rəsmi reyestrində dərc edilmişdir (e-qanun.az/framework/62008).",
      "The decree sets out additional state-support measures to stimulate the export of goods produced outside the oil-and-gas sector. The measures continue the economic-diversification policy of promoting exports, expanding non-oil entrepreneurship and broadening export markets. The full text is published in the official e-Qanun register (e-qanun.az/framework/62008).")

entry("Order", "e-Qanun id 62010", "17 Jun 2026",
      "Məhəmməd Süleyman Əl-Casirin “Dostluq” ordeni ilə təltif edilməsi haqqında",
      "On Awarding Mohammed bin Sulaiman Al-Jasser the 'Dostluq' (Friendship) Order",
      "Sərəncamla Azərbaycan–İslam İnkişaf Bankı əməkdaşlığının inkişafındakı xidmətlərinə görə Məhəmməd Süleyman Əl-Casir “Dostluq” ordeni ilə təltif edilir. Təltif ikitərəfli və çoxtərəfli iqtisadi əməkdaşlığa verilən əhəmiyyəti əks etdirir. (e-qanun.az/framework/62010).",
      "The order confers the 'Dostluq' (Friendship) Order on Mohammed bin Sulaiman Al-Jasser — associated with the Islamic Development Bank — for services to bilateral and multilateral economic cooperation with Azerbaijan. The award reflects the importance attached to that partnership. (e-qanun.az/framework/62010).")

entry("Order", "e-Qanun id 62009", "17 Jun 2026",
      "Azərbaycan Respublikasının səhiyyə işçilərinin təltif edilməsi haqqında",
      "On Awarding Healthcare Workers of the Republic of Azerbaijan",
      "Sərəncamla bir qrup səhiyyə işçisi peşəkar fəaliyyətlərinə və xidmətlərinə görə dövlət təltifləri ilə mükafatlandırılır. Təltif ənənəvi olaraq Tibb İşçiləri gününə təsadüf edir və səhiyyə sahəsinin əməyinin dövlət tərəfindən qiymətləndirilməsini ifadə edir. (e-qanun.az/framework/62009).",
      "The order grants state awards to a group of healthcare workers in recognition of their professional service. The awards traditionally coincide with Medical Workers' Day and express the state's recognition of the health sector. (e-qanun.az/framework/62009).")

entry("Decree", "", "19 Jun 2026, 14:37",
      "Sənaye və kənd təsərrüfatı üzrə tikinti obyektlərinin elektrik şəbəkəsinə qoşulma xidmətinə görə ödənişin subsidiyalaşdırılması haqqında",
      "On Subsidising the Fee for Connecting Industrial and Agricultural Construction Facilities to the Electricity Grid",
      "Fərman sənaye və kənd təsərrüfatı təyinatlı tikinti obyektlərinin elektrik şəbəkəsinə qoşulması üçün ödənilən haqqın dövlət hesabına qismən və ya tam subsidiyalaşdırılmasını nəzərdə tutur. Tədbir investisiya layihələrinin başlanğıc xərclərini azaltmaq, regionlarda emal və istehsal müəssisələrinin yaradılmasını stimullaşdırmaq məqsədi daşıyır. Subsidiyanın həcmi, şərtləri və müraciət qaydası Nazirlər Kabineti tərəfindən müəyyən ediləcəkdir.",
      "The decree provides for partial or full state subsidisation of the fee paid to connect industrial- and agricultural-purpose construction facilities to the electricity grid. The measure is designed to reduce the up-front cost of investment projects and to encourage the establishment of processing and manufacturing enterprises in the regions. The amount, conditions and application procedure for the subsidy are to be set by the Cabinet of Ministers.")

entry("Decree", "", "19 Jun 2026, 14:31",
      "Məktəbəqədər təhsil sahəsində valideyn–dövlət–özəl əməkdaşlığına (PPP) əsaslanan maliyyələşmə modelinin tətbiq edilməsi haqqında",
      "On Introducing a Parent–State–Private Partnership (PPP) Financing Model in Preschool Education",
      "Sənəd məktəbəqədər təhsilin maliyyələşdirilməsində valideyn, dövlət və özəl sektorun birgə iştirakına əsaslanan yeni modelin tətbiqini müəyyən edir. Model bağça xidmətlərinə əlçatanlığı artırmaq, özəl məktəbəqədər müəssisələrin şəbəkəsini genişləndirmək və dövlət subsidiyalarının daha ünvanlı paylanmasını təmin etmək məqsədi daşıyır. Pilot tətbiq, maliyyələşmə normativləri və tərəflərin öhdəlikləri müvafiq icra hakimiyyəti orqanı tərəfindən tənzimlənəcəkdir.",
      "The instrument establishes a new model for funding preschool education based on the joint participation of parents, the state and the private sector. It aims to widen access to nursery services, expand the network of private preschool institutions, and ensure a more targeted allocation of state subsidies. Pilot implementation, funding norms and the obligations of the parties will be regulated by the relevant executive authority.")

entry("Decree", "", "19 Jun 2026, 13:09",
      "Azərbaycan Respublikasının Xarici İşlər Nazirliyində xidmət keçən diplomatik xidmət əməkdaşlarının aylıq vəzifə maaşlarının təsdiq edilməsi haqqında",
      "On Approving the Monthly Position Salaries of Diplomatic-Service Officers Serving in the Ministry of Foreign Affairs",
      "Fərmanla Xarici İşlər Nazirliyinin diplomatik xidmət əməkdaşları üçün yenilənmiş aylıq vəzifə maaşları cədvəli təsdiq olunur. Tədbir diplomatik xidmətdə əmək haqqının müasirləşdirilməsinə və ixtisaslı kadr potensialının saxlanılmasına yönəlib. Maaşlar müvafiq vəzifə dərəcələrinə uyğun olaraq diferensiallaşdırılır.",
      "The decree approves an updated schedule of monthly position salaries for the diplomatic-service officers of the Ministry of Foreign Affairs. The measure is aimed at modernising remuneration in the diplomatic service and retaining qualified personnel. Salaries are differentiated according to the respective grades of office.")

entry("Decree", "", "19 Jun 2026, 13:08",
      "“Diplomatik xidmət haqqında” Azərbaycan Respublikasının Qanununun tətbiqi ilə əlaqədar normativ hüquqi aktlarda dəyişiklik edilməsi haqqında",
      "On Amendments to Normative Legal Acts in Connection with the Application of the Law 'On Diplomatic Service'",
      "Sənəd “Diplomatik xidmət haqqında” Qanunun icrası ilə əlaqədar mövcud normativ bazada uyğunlaşdırıcı dəyişiklikləri əhatə edir. Dəyişikliklər diplomatik rütbələr, vəzifəyə təyinat və xidmət keçmə şərtləri ilə bağlı müddəaları aktuallaşdırır. Eyni gündə imzalanan maaş Fərmanı ilə birlikdə diplomatik xidmət islahatının vahid paketini formalaşdırır.",
      "The instrument introduces conforming amendments to the existing regulatory framework in connection with the implementation of the Law 'On Diplomatic Service.' The changes update provisions on diplomatic ranks, appointments and conditions of service. Together with the salary decree signed the same day, it forms a single package of diplomatic-service reform.")

entry("Decree", "", "19 Jun 2026, 13:06",
      "İnzibati Xətalar Məcəlləsində dəyişikliklərin tətbiqi ilə bağlı tədbirlər haqqında",
      "On Measures Relating to the Application of Amendments to the Code of Administrative Offences",
      "Fərman İnzibati Xətalar Məcəlləsinə edilən dəyişikliklərin tətbiqini təmin edən tədbirləri müəyyən edir. Dəyişikliklər ayrı-ayrı xətaların tərkibinə, sanksiyaların həcminə və ya icraat qaydalarına toxuna bilər. Dəqiq tərkib üçün e-qanun.az-da dərc olunacaq tam mətnə istinad edilməsi tövsiyə olunur.",
      "The decree sets out measures ensuring the application of amendments to the Code of Administrative Offences. The amendments may affect the elements of particular offences, the size of sanctions, or procedural rules. For the precise content, readers are advised to consult the full text to be published on e-qanun.az.")

entry("Decree", "", "19 Jun 2026, 13:05",
      "Azərbaycan Respublikasının Vergi Məcəlləsində və investisiyaların təşviqi ilə bağlı qanunvericilikdə dəyişiklik edilməsi haqqında",
      "On Amendments to the Tax Code and to Legislation on the Promotion of Investments",
      "Sənəd Vergi Məcəlləsinə və investisiyaların təşviqinə dair qanunvericiliyə dəyişiklikləri əhatə edir. Dəyişikliklərin investisiya təşviqi sənədi, güzəştlər və vergi rejimləri ilə bağlı müddəaları aktuallaşdırması gözlənilir. Bu, biznes və vergi məsləhəti üzrə komandalar üçün həftənin ən diqqətçəkən sənədlərindəndir.",
      "The instrument introduces amendments to the Tax Code and to legislation on the promotion of investments. The changes are expected to update provisions on investment-promotion certificates, incentives and tax regimes. This is among the week's most noteworthy instruments for business- and tax-advisory teams.")

entry("Order", "", "19 Jun 2026, 13:04",
      "Azərbaycan Respublikasının Portuqaliya Respublikasında (Lissabon şəhərində) Səfirliyinin fəaliyyətinin təmin edilməsi haqqında",
      "On Ensuring the Operation of the Embassy of the Republic of Azerbaijan in the Portuguese Republic (Lisbon)",
      "Sərəncam Azərbaycanın Lissabon şəhərində fəaliyyət göstərəcək Səfirliyinin maddi-texniki və maliyyə təminatı ilə bağlı tədbirləri müəyyən edir. Sənəd il ərzində Milli Məclisdə müzakirə edilən Lissabonda səfirliyin açılması təşəbbüsünün icra mərhələsini əks etdirir. Tədbir Azərbaycan–Portuqaliya ikitərəfli münasibətlərinin genişlənməsinə xidmət edir.",
      "The order sets out measures for the logistical and financial support of Azerbaijan's Embassy in Lisbon. It reflects the implementation stage of the initiative — discussed in the Milli Majlis earlier this year — to open a resident mission in the Portuguese capital. The measure serves the broadening of bilateral Azerbaijan–Portugal relations.")

notice('info',
       "Qeyd: Bu həftə həmçinin Şuşada keçirilən beynəlxalq konfransa (15 iyun) və Bakı Ombudsmanlar Sammitinə (18 iyun) Prezidentin müraciətləri dərc edilmişdir; bunlar qanunvericilik aktı deyil, rəsmi çıxışlardır.",
       "Note: the President's addresses to the international conference in Shusha (15 June) and the Baku Ombudsmen Summit (18 June) were also published this week; these are official statements rather than legislative acts.")

# ── SECTION 2: CABINET ────────────────────────────────────────────────────────
section_header("Nazirlər Kabineti · Cabinet of Ministers",
               "Cabinet of Ministers", "Nazirlər Kabineti", "Decisions & Orders — 16–19 June 2026")
lead("Hesabat dövründə Nazirlər Kabineti altı Qərar (№181–186) və dörd icra Sərəncamı (№408s, 409s, 419s, 420s) qəbul etmişdir. Sərəncamlar Prezidentin əvvəlki həftələrdə imzaladığı 676, 679, 682 və 684 nömrəli Fərmanların icrasını təmin edir.",
     "During the reporting period the Cabinet adopted six Decisions (Nos. 181–186) and four implementing Orders (Nos. 408s, 409s, 419s, 420s). The Orders give effect to Presidential Decrees 676, 679, 682 and 684 signed in the preceding weeks.")

decisions = [
 ("Decision", "No. 181 · e-Qanun id 62027", "16 Jun 2026",
  "“Elektrik şəbəkəsinə dair Qaydalar”da (Nazirlər Kabinetinin 2024-cü il 25 iyun tarixli 315 nömrəli Qərarı) dəyişiklik edilməsi haqqında",
  "On Amendments to the 'Rules on the Electricity Grid' (CoM Decision No. 315 of 25 June 2024)",
  "Qərar 2024-cü ildə təsdiq edilmiş elektrik şəbəkəsi qaydalarına dəyişiklik edir; şəbəkəyə qoşulma, texniki tələblər və ya tənzimləmə prosedurları yenilənir. Dəyişiklik eyni həftə imzalanan elektrik qoşulması subsidiyası Fərmanı ilə uzlaşır. Tam mətn: e-qanun.az/framework/62027.",
  "The decision amends the 2024 electricity-grid rules, updating grid-connection, technical-requirement or regulatory procedures. The change dovetails with the grid-connection subsidy decree signed the same week. Full text: e-qanun.az/framework/62027."),
 ("Decision", "No. 182 · e-Qanun id 62025", "16 Jun 2026",
  "“Dövlət Xidmətlərinin Elektron Reyestrinin aparılması Qaydaları”nın (2015-ci il 9 fevral tarixli 32 nömrəli Qərar) və “Elektron xidmətlərin təşkili və göstərilməsi qaydaları”nın (2025-ci il 4 avqust tarixli 241 nömrəli Qərar) dəyişdirilməsi barədə",
  "On Amendments to the 'Rules for Maintaining the Electronic Registry of State Services' (CoM Decision No. 32 of 9 Feb 2015) and the 'Rules for Organising and Providing Electronic Services' (CoM Decision No. 241 of 4 Aug 2025)",
  "Qərar həm dövlət xidmətlərinin elektron reyestrinin aparılması qaydalarını, həm də 2025-ci il avqustda təsdiq edilmiş elektron xidmətlərin təşkili qaydalarını yeniləyir. Tədbir rəqəmsal dövlət xidmətlərinin idarə olunmasının təkmilləşdirilməsinə yönəlib. Tam mətn: e-qanun.az/framework/62025.",
  "The decision updates both the rules for maintaining the electronic registry of state services and the August-2025 rules for organising electronic services. The measure improves the management of digital public services. Full text: e-qanun.az/framework/62025."),
 ("Decision", "No. 183 · e-Qanun id 62026", "16 Jun 2026",
  "“Dəmiryol nəqliyyatında mülki müdafiənin təşkili Qaydaları”nın təsdiq edilməsi haqqında",
  "On Approval of the 'Rules for the Organisation of Civil Defence in Railway Transport'",
  "Qərarla dəmiryol nəqliyyatı sahəsində mülki müdafiənin təşkilinə dair yeni Qaydalar təsdiq edilir. Sənəd fövqəladə hallarda hazırlıq, xəbərdarlıq və mühafizə tədbirlərini tənzimləyir. Tam mətn: e-qanun.az/framework/62026.",
  "The decision approves new Rules for organising civil defence in the railway-transport sector. The instrument governs emergency preparedness, warning and protection measures. Full text: e-qanun.az/framework/62026."),
 ("Decision", "No. 184", "19 Jun 2026",
  "Azərbaycan Dövlət Su Ehtiyatları Agentliyinin tabeliyində olan qapalı səhmdar cəmiyyətlərinin nizamnamələrinin təsdiq edilməsi haqqında",
  "On Approval of the Charters of the Closed Joint-Stock Companies Subordinate to the Azerbaijan State Water Resources Agency",
  "Qərar Su Ehtiyatları Agentliyinin tabeliyindəki qapalı səhmdar cəmiyyətlərinin nizamnamələrini təsdiq edir; idarəetmə strukturu, səlahiyyətlər və fəaliyyət predmeti müəyyənləşdirilir. Tədbir su təsərrüfatının korporativ idarəetməsinin qaydaya salınmasına xidmət edir.",
  "The decision approves the charters of the closed joint-stock companies subordinate to the State Water Resources Agency, defining their governance structure, powers and objects of activity. The measure serves to regularise corporate governance in the water sector."),
 ("Decision", "No. 185", "19 Jun 2026",
  "“Enerji balansının hazırlanması, təsdiq edilməsi və icrasına nəzarət Qaydası”nda (2024-cü il 1 may tarixli 233 nömrəli Qərar) dəyişiklik edilməsi haqqında",
  "On Amendments to the 'Rules for the Preparation, Approval and Monitoring of the Energy Balance' (CoM Decision No. 233 of 1 May 2024)",
  "Qərar enerji balansının hazırlanması, təsdiqi və icrasına nəzarət qaydasına dəyişiklik edir. Dəyişiklik enerji planlaması və hesabatlılıq mexanizmlərinin dəqiqləşdirilməsinə yönəlib.",
  "The decision amends the rules for preparing, approving and monitoring the energy balance. The change is aimed at refining energy-planning and reporting mechanisms."),
 ("Decision", "No. 186", "19 Jun 2026",
  "Bərpa olunan enerji mənbələrinin ərazisində günəş elektrik stansiyasının tikintisinə icazə alınması müddətinin uzadılması haqqında",
  "On Extending the Period for Obtaining a Permit for the Construction of a Solar Power Plant in a Renewable Energy Zone",
  "Qərar bərpa olunan enerji zonasında günəş elektrik stansiyasının tikintisinə icazənin alınması müddətini uzadır. Tədbir investorlara layihələrin reallaşdırılması üçün əlavə vaxt verir və yaşıl enerji siyasətinin davamı kimi qiymətləndirilir.",
  "The decision extends the deadline for obtaining a construction permit for a solar power plant in a renewable-energy zone. The measure gives investors additional time to realise projects and is read as a continuation of the green-energy policy."),
]
for d in decisions:
    entry(*d)

orders = [
 ("Order", "No. 408s", "16 Jun 2026",
  "Prezidentin 2026-cı il 2 iyun tarixli 676 nömrəli Fərmanının (kibertəhlükəsizlik sahəsində idarəetmənin təkmilləşdirilməsi) icrasının təmin edilməsi barədə",
  "On Ensuring Implementation of Presidential Decree No. 676 of 2 June 2026 (Improvement of Cybersecurity Governance)",
  "Sərəncam kibertəhlükəsizlik sahəsində idarəetmənin təkmilləşdirilməsinə dair 676 nömrəli Prezident Fərmanının icra tədbirlərini müəyyən edir. İcra orqanlarına konkret tapşırıqlar verilir və məsul qurumlar göstərilir.",
  "The order sets out the implementation measures for Presidential Decree No. 676 on improving cybersecurity governance, assigning concrete tasks to executive bodies and designating responsible agencies."),
 ("Order", "No. 409s", "16 Jun 2026",
  "Prezidentin 2026-cı il 8 iyun tarixli 679 nömrəli Fərmanının (Prezident fərmanlarında dəyişikliklər) icrasının təmin edilməsi haqqında",
  "On Ensuring Implementation of Presidential Decree No. 679 of 8 June 2026 (Amendments to Presidential Decrees)",
  "Sərəncam 679 nömrəli Fərmanın icrasını təmin edir; mənbə məlumatına görə dəyişikliklər kənd təsərrüfatı subsidiyaları və elektron kənd təsərrüfatı sistemləri ilə bağlı aktları əhatə edir. Konkret tədbirlər müvafiq nazirliklərə həvalə olunur.",
  "The order gives effect to Decree No. 679; per the source, the underlying amendments concern acts on agricultural subsidies and the electronic-agriculture system. Specific measures are entrusted to the relevant ministries."),
 ("Order", "No. 419s", "19 Jun 2026",
  "Prezidentin 2026-cı il 12 iyun tarixli 684 nömrəli Fərmanının (“Hava limanlarının (aeroportlarının) idarə edilməsi Qaydası”) icrasının təmin edilməsi barədə",
  "On Ensuring Implementation of Presidential Decree No. 684 of 12 June 2026 ('Rules for the Administration of Airports')",
  "Sərəncam hava limanlarının (aeroportlarının) idarə edilməsi qaydasını təsdiq edən 684 nömrəli Fərmanın icra tədbirlərini müəyyən edir. Sənəd aviasiya infrastrukturunun idarəolunması və tənzimlənməsi baxımından əhəmiyyət daşıyır.",
  "The order establishes the implementation measures for Decree No. 684, which approves the rules for administering airports. It is significant for the management and regulation of aviation infrastructure."),
 ("Order", "No. 420s", "19 Jun 2026",
  "Prezidentin 2026-cı il 10 iyun tarixli 682 nömrəli Fərmanının (“Vətəndaşların müraciətləri haqqında” Qanunda dəyişiklik) icrasının təmin edilməsi barədə",
  "On Ensuring Implementation of Presidential Decree No. 682 of 10 June 2026 (Amendment to the Law 'On Citizens' Appeals')",
  "Sərəncam “Vətəndaşların müraciətləri haqqında” Qanunda dəyişiklik edən 682 nömrəli Fərmanın icrasını təmin edir; dövlət və bələdiyyə orqanlarında müraciətlərə baxılma proseduru yenilənir. Tədbir vətəndaş–dövlət münasibətlərində şəffaflığı artırmağa yönəlib.",
  "The order gives effect to Decree No. 682 amending the Law 'On Citizens' Appeals,' updating the procedure for handling appeals in state and municipal bodies. The measure aims to increase transparency in citizen–state interaction."),
]
for o in orders:
    entry(*o)

# ── SECTION 3: MILLI MAJLIS ───────────────────────────────────────────────────
section_header("Azərbaycan Respublikasının Milli Məclisi · Parliament",
               "Milli Majlis", "Milli Məclis", "Bills Tabled — 19 June 2026")
lead("Milli Məclis 19 iyun 2026-cı il tarixli plenar iclasında on qanun layihəsini müzakirəyə çıxarmışdır. Hesabat dövründə yeni qanun və ya parlament qərarı qəbul edilməmişdir (ən son Qanun — №415-VIIQD, 26 may 2026). Layihələr müzakirə mərhələsindədir.",
     "At its plenary sitting of 19 June 2026 the Milli Majlis tabled ten bills for discussion. No new law or parliamentary decree was enacted in the reporting period (the most recent Law — No. 415-VIIQD — dates to 26 May 2026). The bills are at the discussion stage.")

bills = [
 ("Azərbaycan Respublikasının 2025-ci il dövlət büdcəsinin icrası haqqında",
  "On the Execution of the 2025 State Budget of the Republic of Azerbaijan",
  "İllik hesabatlılıq sənədi — 2025-ci il dövlət büdcəsinin icrasının parlament tərəfindən təsdiqi.",
  "Annual accountability instrument — parliamentary approval of the execution of the 2025 state budget."),
 ("Azərbaycan Respublikasının Gömrük Məcəlləsində dəyişiklik edilməsi haqqında",
  "On Amendments to the Customs Code",
  "Gömrük prosedurları, rüsumlar və ya gömrük nəzarəti ilə bağlı müddəaların yenilənməsi.",
  "Updates to provisions on customs procedures, duties or customs control."),
 ("“Dövlət qulluğu haqqında” Azərbaycan Respublikasının Qanununda dəyişiklik edilməsi barədə",
  "On Amendments to the Law 'On Civil Service'",
  "Dövlət qulluqçularının hüquq, vəzifə və ya karyera tənzimləməsinə dair dəyişikliklər.",
  "Amendments concerning the rights, duties or career regulation of civil servants."),
 ("Azərbaycan–Ekvador arasında diplomatik, rəsmi və xidməti pasport sahiblərinin viza tələblərindən azad edilməsi Sazişinin təsdiqi",
  "Ratification of the Azerbaijan–Ecuador Agreement on Visa Exemption for Holders of Diplomatic, Official and Service Passports",
  "İki ölkə arasında rəsmi pasport sahibləri üçün vizasız rejim yaradan beynəlxalq sazişin ratifikasiyası.",
  "Ratification of an international agreement establishing visa-free travel for official-passport holders of the two states."),
 ("“Normativ hüquqi aktlar haqqında” Azərbaycan Respublikasının Konstitusiya Qanununda dəyişiklik edilməsi barədə",
  "On Amendments to the Constitutional Law 'On Normative Legal Acts'",
  "Qanunvericilik texnikası və normativ aktların qəbulu prosesinə dair konstitusiya səviyyəli dəyişikliklər.",
  "Constitutional-level amendments to legislative drafting technique and the process of adopting normative acts."),
 ("“Dövlət qulluğu haqqında” və “Media haqqında” Azərbaycan Respublikasının qanunlarında dəyişiklik edilməsi barədə",
  "On Amendments to the Laws 'On Civil Service' and 'On Media'",
  "İki qanunu eyni zamanda dəyişdirən paket — dövlət qulluğu və media tənzimləməsi sahələrində.",
  "A package amending two laws simultaneously — in the fields of civil service and media regulation."),
 ("Ticarət Gəmiçiliyi Məcəlləsi, “Yol hərəkəti”, “Nəqliyyat”, “Dövlət rüsumu”, məhdud dövriyyəli əşyalar, “Avtomobil nəqliyyatı” və “Aviasiya” qanunlarında dəyişiklik (omnibus)",
  "Omnibus Amendments to the Merchant Shipping Code and the Laws on Road Traffic, Transport, State Duty, Restricted-Circulation Items, Road Transport and Aviation",
  "Nəqliyyatın bütün növlərini (dəniz, avtomobil, aviasiya) əhatə edən geniş çoxqanunlu islahat paketi.",
  "A broad multi-law reform package covering all modes of transport — maritime, road and aviation."),
 ("“Yol hərəkəti haqqında” Qanunun tətbiqi ilə əlaqədar 307-VIIQD nömrəli Qanunun (9.12.2025) icrası üçün bəzi qanunlarda dəyişiklik",
  "Amendments to Certain Laws to Implement Law No. 307-VIIQD (9 Dec 2025) Amending the Law 'On Road Traffic'",
  "2025-ci ilin sonunda qəbul edilmiş yol hərəkəti islahatının icrasını təmin edən əlaqədar dəyişikliklər.",
  "Conforming amendments giving effect to the road-traffic reform enacted in late 2025."),
 ("Milli Məclisin Daxili Nizamnaməsi, “İctimai iştirakçılıq” və vətəndaşların qanunvericilik təşəbbüsü hüququ haqqında qanunlarda dəyişiklik",
  "On Amendments to the Laws on the Milli Majlis Internal Charter, 'On Public Participation' and on Citizens' Right of Legislative Initiative",
  "İctimai dinləmələrin və qanunvericilik təşəbbüsünün Elektron Hökumət sistemi üzərindən təşkilinə imkan verən islahat.",
  "A reform enabling public hearings and citizens' legislative initiative to be organised via the e-Government system."),
 ("“Ümumi təhsil haqqında” Azərbaycan Respublikasının Qanununda dəyişiklik edilməsi barədə",
  "On Amendments to the Law 'On General Education'",
  "Ümumi təhsilin təşkili, idarəolunması və ya standartlarına dair dəyişikliklər.",
  "Amendments concerning the organisation, governance or standards of general education."),
]
for i, (az_t, en_t, az_b, en_b) in enumerate(bills, 1):
    entry("Bill", f"Layihə {i}/10", "19 Jun 2026 · Müzakirədə / Under discussion", az_t, en_t, az_b, en_b)

# ── SECTION 4: CONSTITUTIONAL COURT ───────────────────────────────────────────
section_header("Konstitusiya Məhkəməsi · Constitutional Court",
               "Constitutional Court", "Konstitusiya Məhkəməsi", "constcourt.gov.az")
notice('info',
       "Bu həftə yeni qərar YOXDUR. Konstitusiya Məhkəməsinin öz portalı tərtibat zamanı (22 iyun) HTTP 503 xətası versə də, e-Qanun rəsmi reyestri üzrə çarpaz yoxlama göstərir ki, Məhkəmənin ən son qərarı 25.05.2026 tarixlidir — 15–22 iyun dövründə yeni qərar qəbul edilməmişdir.",
       "No new decisions this week. Although the Court's own portal returned HTTP 503 at compilation (22 June), a cross-check against the official e-Qanun register confirms the Court's most recent decision is dated 25.05.2026 — none were issued in the 15–22 June window.")

# ── SECTION 5: E-QANUN ────────────────────────────────────────────────────────
section_header("Rəsmi Qanunvericilik Reyestri · Official Legislative Register",
               "e-Qanun", "e-Qanun", "api.e-qanun.az · 15–22 June 2026")
lead("Aşağıdakı sənədlər tərtibat tarixinə (22 iyun) e-Qanun rəsmi reyestrində qeydiyyata alınmış və bu həftəyə (qəbul tarixi 15–22 iyun) aiddir. Qeyd: 19 iyun tarixli Prezident Fərmanları/Sərəncamı və Nazirlər Kabinetinin 184–186 saylı qərarları ilə icra sərəncamları hələ reyestrə daxil edilməmişdir — adətən bir neçə günlük qeydiyyat gecikməsi olur.",
     "The documents below were registered in the official e-Qanun register as of the compilation date (22 June) and fall within this week (accept date 15–22 June). Note: the 19 June presidential instruments and Cabinet decisions 184–186 with their implementing orders had not yet been entered in the register — a registration lag of a few days is normal.")

reg_rows = [
 ("Decree",   "16 Jun 2026", "62008", "Qeyri-neft-qaz malları üzrə ixracın dəstəklənməsi ilə bağlı əlavə tədbirlər", "Non-oil-and-gas export support — additional measures"),
 ("Order",    "17 Jun 2026", "62010", "Məhəmməd Süleyman Əl-Casirin “Dostluq” ordeni ilə təltif edilməsi", "Award of 'Dostluq' Order to Mohammed Al-Jasser"),
 ("Order",    "17 Jun 2026", "62009", "Azərbaycan Respublikasının səhiyyə işçilərinin təltif edilməsi", "Award of state honours to healthcare workers"),
 ("Decision", "16 Jun 2026", "62027", "Elektrik şəbəkəsinə dair Qaydalarda dəyişiklik (NK Qərarı №181)", "Electricity-grid rules amendment (CoM Decision No. 181)"),
 ("Decision", "16 Jun 2026", "62025", "Dövlət xidmətlərinin elektron reyestri və e-xidmət qaydalarında dəyişiklik (№182)", "Electronic registry & e-services rules amendment (No. 182)"),
 ("Decision", "16 Jun 2026", "62026", "Dəmiryol nəqliyyatında mülki müdafiə Qaydaları (№183)", "Railway-transport civil-defence rules (No. 183)"),
]
rt = doc.add_table(rows=1, cols=4)
rt.style = 'Table Grid'
hdr = rt.rows[0].cells
for c, label in zip(hdr, ["Type", "Date", "Framework", "Document (AZ) / Translation (EN)"]):
    p = c.paragraphs[0]; no_space(p); shade_cell(c, '1a1a1a')
    run(p, label, size=7.5, bold=True, color=WHITE, font=SANS, caps=True)
for typ, dt, fid, az, en in reg_rows:
    cells = rt.add_row().cells
    p0 = cells[0].paragraphs[0]; no_space(p0); run(p0, typ, size=8, bold=True, font=SANS)
    p1 = cells[1].paragraphs[0]; no_space(p1); run(p1, dt, size=8, font=SANS)
    p2 = cells[2].paragraphs[0]; no_space(p2)
    run(p2, "/framework/" + fid, size=8, color=RGBColor(0x1a,0x3a,0x5a), font=SANS)
    p3 = cells[3].paragraphs[0]; no_space(p3, 0, 1, line=1.2); run(p3, az, size=8.5, bold=True)
    p3b = cells[3].add_paragraph(); no_space(p3b, 0, 0, line=1.2); run(p3b, en, size=8, italic=True, color=GREY)

m = doc.add_paragraph(); no_space(m, 4, 4, line=1.3)
run(m, "Metod: e-qanun.az JavaScript ilə işləyən tətbiqdir; məlumat api.e-qanun.az daxili API-si (getDetailSearch, reyestr tipi üzrə) vasitəsilə birbaşa əldə edilmişdir. ",
    size=8, italic=True, color=LGREY, font=SANS)
run(m, "Method: e-qanun.az is a JavaScript app; data was retrieved directly from its api.e-qanun.az backend (getDetailSearch, scoped per register type). Full texts at the framework links above.",
    size=8, italic=True, color=LGREY, font=SANS)

# ── REGULATORY SPOTLIGHT ──────────────────────────────────────────────────────
section_header("Təhlil · Analysis", "Regulatory Spotlight", "Tənzimləyici Diqqət", "Week 25 / 2026")
spot = [
 ("1. Diplomatik xidmət islahatı / Diplomatic-service reform",
  "Eyni gündə imzalanan üç sənəd — XİN maaşları Fərmanı, “Diplomatik xidmət haqqında” Qanunun tətbiqi dəyişiklikləri və Lissabon səfirliyi Sərəncamı — diplomatik xidmətin müasirləşdirilməsi üzrə əlaqəli paketi formalaşdırır. Beynəlxalq və dövlət hüququ üzrə komandalar üçün aktualdır.",
  "Three same-day instruments — the MFA salary decree, the diplomatic-service-law application amendments, and the Lisbon embassy order — form a coordinated package modernising the diplomatic service. Relevant for international- and public-law teams."),
 ("2. Vergi və investisiya təşviqi / Tax & investment promotion",
  "Vergi Məcəlləsinə və investisiya təşviqi qanunvericiliyinə dəyişiklik bu həftənin biznes baxımından ən mühüm sənədidir. Tam mətn dərc olunan kimi güzəştlər və investisiya təşviqi sənədi rejimi diqqətlə nəzərdən keçirilməlidir.",
  "The amendments to the Tax Code and investment-promotion legislation are the week's most commercially significant instrument. Once the full text is published, the incentive and investment-certificate regime should be reviewed closely."),
 ("3. Enerji və “yaşıl” siyasət / Energy & green policy",
  "Elektrik qoşulması subsidiyası Fərmanı, elektrik şəbəkəsi (181) və enerji balansı (185) qaydalarına dəyişikliklər, günəş stansiyası icazə müddətinin uzadılması (186) — enerji sahəsində ardıcıl tənzimləyici fəaliyyəti göstərir.",
  "The grid-connection subsidy decree, amendments to the electricity-grid (181) and energy-balance (185) rules, and the solar-permit extension (186) together signal sustained regulatory activity in the energy sector."),
 ("4. Rəqəmsal idarəetmə və iştirakçılıq / Digital governance & participation",
  "Dövlət Xidmətlərinin Elektron Reyestri (182), kibertəhlükəsizlik idarəetməsi (676/408s) və ictimai iştirakçılığın e-Hökumət üzərindən təşkili layihəsi rəqəmsallaşma istiqamətini gücləndirir. Tənzimlənən sahələrdə uyğunluq komandaları izləməlidir.",
  "The electronic registry of state services (182), cybersecurity governance (676/408s), and the bill to organise public participation via e-Government reinforce the digitalisation agenda. Compliance teams in regulated sectors should monitor it."),
]
for title_, az_, en_ in spot:
    h = doc.add_paragraph(); no_space(h, 6, 1)
    run(h, title_, size=9.5, bold=True)
    a = doc.add_paragraph(); no_space(a, 0, 1, line=1.3)
    run(a, az_, size=8.5, color=GREY)
    e = doc.add_paragraph(); no_space(e, 0, 4, line=1.3)
    run(e, en_, size=8.5, italic=True, color=GREY)

# ── FOOTER ────────────────────────────────────────────────────────────────────
f = doc.add_paragraph(); no_space(f, 14, 2)
set_borders(f, sides=('top',), sz=18)
run(f, "Omni Law Gazette", size=8, bold=True, font=SANS)
run(f, " — Omni Law Firm tərəfindən yalnız məlumat məqsədilə hazırlanmış daxili icmaldır və hüquqi məsləhət təşkil etmir. Sənəd təsvirləri rəsmi dövlət portallarından götürülmüşdür; etibar etməzdən əvvəl ilkin mənbələrə və tam mətnlərə müraciət edin.",
    size=7.5, color=GREY, font=SANS)
f2 = doc.add_paragraph(); no_space(f2, 0, 2)
run(f2, "An internal digest compiled by Omni Law Firm for information only; not legal advice. Descriptions are derived from official government portals — consult primary sources and full texts before relying on any item.",
    size=7.5, italic=True, color=GREY, font=SANS)
f3 = doc.add_paragraph(); no_space(f3, 0, 0)
run(f3, "Sources: president.az · nk.gov.az · meclis.gov.az · e-qanun.az · constcourt.gov.az      |      Issue No. 2 · Week of 15–22 June 2026 · Compiled 22 June 2026 · © 2026 Omni Law Firm",
    size=7.5, color=LGREY, font=SANS)

# ── save ──────────────────────────────────────────────────────────────────────
out_dir = os.path.join(os.path.dirname(__file__), 'gazettes')
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'gazette-2026-06-22.docx')
doc.save(out_path)
print(f"✅ Saved {out_path}")
