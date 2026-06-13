#!/usr/bin/env python3
"""Build an editable Word (.docx) edition of the Omni Law Gazette, Issue No. 2.

Renders the same content as gazette-2026-06-13-issue-2.html as a native,
fully-editable Word document (headings, tables, styled runs) rather than an
HTML import — so colleagues can edit it directly in Word.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

INK   = RGBColor(0x1a, 0x1a, 0x1a)
GREY  = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xff, 0xff, 0xff)
RED   = RGBColor(0x5a, 0x1a, 0x1a)
GREEN = RGBColor(0x1a, 0x4a, 0x1a)
NAVY  = RGBColor(0x2c, 0x3e, 0x50)

doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Georgia'
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(6)

# Narrow margins
for s in doc.sections:
    s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor); tcPr.append(sh)


def hrule(p=None):
    p = p or doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1a1a1a')
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def run(p, text, *, size=10, bold=False, italic=False, color=INK, font='Georgia',
        caps=False, spacing=None):
    r = p.add_run(text.upper() if caps else text)
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold
    r.font.italic = italic; r.font.color.rgb = color
    if spacing is not None:
        rPr = r._r.get_or_add_rPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(spacing)); rPr.append(sp)
    return r


def para(align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


# ── MASTHEAD ────────────────────────────────────────────────────────────────
hrule(para(after=2))
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=0)
run(p, 'OMNI ', size=30, bold=True, font='Georgia')
run(p, 'Law', size=30, italic=True, font='Georgia')
run(p, ' GAZETTE', size=30, bold=True, font='Georgia')
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
run(p, 'L E G I S L A T I V E   I N T E L L I G E N C E   F O R   L E G A L   P R O F E S S I O N A L S',
    size=8, color=GREY, font='Arial')
hrule(para(after=2))
p = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
run(p, 'ISSUE No. 2', size=8, bold=True, font='Arial', color=GREY)
run(p, '      |      REPORTING PERIOD: 8–13 JUNE 2026      |      PUBLISHED: 13 JUNE 2026',
    size=8, font='Arial', color=GREY)

# Week banner (shaded one-cell table)
tb = doc.add_table(rows=1, cols=1); tb.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tb.rows[0].cells[0]; shade(cell, '1a1a1a')
cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(cp, 'WEEK 24, 2026  ·  REPUBLIC OF AZERBAIJAN  ·  LEGISLATIVE MONITOR',
    size=8, bold=True, color=WHITE, font='Arial')
doc.add_paragraph().paragraph_format.space_after = Pt(2)


def section(institution, title, en):
    hrule(para(before=10, after=2))
    p = para(after=2)
    run(p, institution.upper() + '   ', size=7, color=GREY, font='Arial')
    run(p, title, size=15, bold=True)
    p2 = para(after=6)
    run(p2, en, size=8, italic=True, color=GREY, font='Arial')


def subhead(text):
    p = para(before=8, after=4)
    run(p, text.upper(), size=8, bold=True, color=GREY, font='Arial', spacing=20)
    hrule(p)


def doc_item(tag, tagcolor, number, date, title, body, ref=None):
    p = para(before=4, after=1)
    # tag chip
    r = run(p, f' {tag.upper()} ', size=7, bold=True, color=WHITE, font='Arial')
    rPr = r._r.get_or_add_rPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear')
    sh.set(qn('w:fill'), tagcolor); rPr.append(sh)
    if number: run(p, f'  {number}', size=8, color=GREY, font='Arial')
    run(p, f'\t{date}', size=8, color=GREY, font='Arial')
    p2 = para(after=1); run(p2, title, size=9.5, bold=True)
    p3 = para(after=2); run(p3, body, size=9, color=RGBColor(0x33,0x33,0x33))
    if ref:
        p4 = para(after=4); run(p4, ref, size=7.5, italic=True, color=GREY, font='Arial')


# ── AT A GLANCE ─────────────────────────────────────────────────────────────
subhead('This Week at a Glance — 8 to 13 June 2026')
glance = [('6','Presidential Decrees'),('1','Presidential Order'),('11','Cabinet Decisions'),
          ('4','Cabinet Orders'),('10','Bills in Parliament'),('11','Diplomatic Letters')]
gt = doc.add_table(rows=2, cols=6); gt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(n,l) in enumerate(glance):
    c1 = gt.cell(0,i).paragraphs[0]; c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(c1, n, size=20, bold=True, color=NAVY)
    c2 = gt.cell(1,i).paragraphs[0]; c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(c2, l, size=7, color=GREY, font='Arial')

# ── LEAD ────────────────────────────────────────────────────────────────────
p = para(before=8, after=8)
run(p, 'A busy legislative week saw the President of the Republic sign ', size=10)
run(p, 'six decrees', size=10, bold=True)
run(p, ' spanning civil-aviation governance, citizens’ rights, the capital’s administration, and a '
       'pronounced economic-stimulus thrust — new mechanisms for entrepreneur credit guarantees, '
       'non-oil export support, and agrarian finance. The Cabinet of Ministers issued ', size=10)
run(p, 'eleven decisions and four orders', size=10, bold=True)
run(p, ', advancing the renewable-energy land programme and adopting major sectoral rules for railways '
       'and gas supply. Parliament placed a wide-ranging package of bills — headlined by a Land Code '
       'overhaul — before its members on 12 June.', size=10)

# ── PRESIDENTIAL OFFICE ─────────────────────────────────────────────────────
section('Administration of the President', 'Presidential Office',
        'Decrees, Orders & Correspondence — 8–12 June 2026')

subhead('Decrees (Fərmanlar)')
doc_item('Decree','5a2c2c',None,'12 Jun 2026',
    'Approval of the Rules on the Management of Airports (Aerodromes)',
    'Adopts a comprehensive regulatory framework for the management and operation of the country’s '
    'airports, issued under “measures to improve management in the field of civil aviation.” '
    'Establishes governance, operational standards and operator responsibilities.',
    'Constitution, Art. 109(32) · Cabinet directed to implement')
doc_item('Decree','5a2c2c',None,'10 Jun 2026',
    'Application of the Law on Citizens’ Appeals (No. 415-VIIQD) & Amendment to Decree No. 950 (2016)',
    'Brings into force the 26 May 2026 amendments to the Law on Citizens’ Appeals and amends the 2016 '
    'Rules on handling citizens’ appeals in state and municipal bodies, state-controlled legal '
    'entities and budget organisations.',
    'Constitution, Art. 109(19),(32) · Implements Law 415-VIIQD')
doc_item('Decree','5a2c2c',None,'10 Jun 2026',
    'Improvement of the Structure and Management of the Baku City Executive Authority',
    'Restructures the administration of the capital, adjusting the organisational framework and '
    'management arrangements of the Baku City Executive Authority.',
    'president.az/az/articles/view/72724 (full text + PDF)')
doc_item('Decree','5a2c2c',None,'9 Jun 2026',
    'Additional Measures to Improve Guarantee & Interest-Subsidy Mechanisms for Credits to Entrepreneurs',
    'Enhances the State’s instruments for guaranteeing, and subsidising the interest on, loans '
    'extended to business entities, and amends a series of related presidential decrees.',
    'Amends several earlier presidential decrees')
doc_item('Decree','5a2c2c',None,'9 Jun 2026',
    'Additional Measures to Stimulate Non-Oil-and-Gas Exports & Support Transport Costs',
    'Introduces further incentives for exporters of non-oil-and-gas goods, including support for '
    'transport/logistics costs, in furtherance of economic diversification.',
    'Economic diversification / non-oil sector')
doc_item('Decree','5a2c2c',None,'8 Jun 2026',
    'Amendments to Decrees on the Agrarian Credit & Development Agency and State Support to Agriculture',
    'Amends the 2015 statute (Decree No. 571) of the Agrarian Credit and Development Agency under the '
    'Ministry of Agriculture and related instruments governing state support to agriculture.',
    'Amends Decree No. 571 (23 Jul 2015)')

subhead('Orders (Sərəncamlar)')
doc_item('Order','4a4a4a',None,'8 Jun 2026',
    'Award of the “Honorary Diploma of the President” to R. M. Fataliyev',
    'Confers the Presidential Honorary Diploma in recognition of services, under the President’s '
    'constitutional honours power.',
    'Constitution, Art. 109(23)')

p = para(before=4, after=6)
run(p, 'Editor’s note: ', size=8.5, bold=True, color=GREY, font='Arial')
run(p, 'president.az publishes decree texts and dates but not the sequential document number on the HTML '
       'view; numbers are assigned on publication in the consolidated register (e-qanun.az). Confirm before citation.',
    size=8.5, color=GREY, font='Arial')

subhead('Diplomatic Correspondence (Məktublar)')
p = para(after=4)
run(p, 'Eleven letters were exchanged, clustered around Azerbaijan’s National Salvation Day (15 June). '
       '“IN” = received by the President of Azerbaijan; “OUT” = sent by him.', size=9, font='Arial', color=GREY)
letters = [
    ('12 Jun','IN','Fuat Oktay','Chair, Foreign Relations Cttee., GNA of Türkiye','National Salvation Day greeting'),
    ('12 Jun','IN','Benjamin Netanyahu','Prime Minister of Israel','National Salvation Day greeting'),
    ('12 Jun','OUT','Ferdinand R. Marcos','President of the Philippines','Bilateral / Independence Day'),
    ('12 Jun','OUT','Vladimir Putin','President of the Russian Federation','Bilateral relations'),
    ('11 Jun','IN','Edi Rama','Prime Minister of Albania','National Salvation Day greeting'),
    ('11 Jun','IN','Adama Barrow','President of The Gambia','National Salvation Day greeting'),
    ('11 Jun','IN','Haji Hassanal Bolkiah','Sultan of Brunei Darussalam','National Salvation Day greeting'),
    ('10 Jun','IN','Bajram Begaj','President of Albania','National Salvation Day greeting'),
    ('10 Jun','OUT','António José Seguro','President of Portugal','Bilateral relations'),
    ('9 Jun','IN','Nataša Pirc Musar','President of Slovenia','National Salvation Day greeting'),
    ('8 Jun','IN','José Raúl Mulino Quintero','President of Panama','National Salvation Day greeting'),
]
lt = doc.add_table(rows=1, cols=5); lt.style = 'Table Grid'
hdr = ['Date','Dir.','Counterpart','Country / Office','Context']
for i,h in enumerate(hdr):
    c = lt.rows[0].cells[i]; shade(c,'1a1a1a'); cp = c.paragraphs[0]
    run(cp, h, size=7.5, bold=True, color=WHITE, font='Arial')
for d,dir_,party,country,ctx in letters:
    cells = lt.add_row().cells
    run(cells[0].paragraphs[0], d, size=8.5, font='Arial')
    run(cells[1].paragraphs[0], dir_, size=8.5, bold=True, font='Arial',
        color=GREEN if dir_=='IN' else RED)
    run(cells[2].paragraphs[0], party, size=8.5)
    run(cells[3].paragraphs[0], country, size=8.5)
    run(cells[4].paragraphs[0], ctx, size=8.5, color=GREY)

# ── CABINET OF MINISTERS ────────────────────────────────────────────────────
section('Nazirlər Kabineti', 'Cabinet of Ministers', 'Decisions & Orders — 9–12 June 2026')
p = para(after=4)
run(p, 'Eleven Decisions and four Orders, all dated 9, 10 and 12 June (no instruments dated 8, 11 or 13 June).',
    size=9, font='Arial', color=GREY)

decisions = [
    ('178','12 Jun 2026','Renewable Energy Land Designation — Hacıqabul District (Parcel I)',
     'Designates a state-owned land parcel in Hacıqabul district as a renewable energy source area.'),
    ('177','12 Jun 2026','Renewable Energy Land Designation — Hacıqabul District (Parcel II)',
     'Designates a second state-owned land parcel in Hacıqabul district as a renewable energy source area.'),
    ('176','12 Jun 2026','Amendment — Jabrayil District Renewable Energy Land (Decision No. 466 of 2024)',
     'Amends CoM Decision No. 466 (29 Oct 2024) on designating state land in Jabrayil district as a renewable energy area.'),
    ('175','10 Jun 2026','Amendment to the List of Entities Outside the Ministry of Economy Structure',
     'Amends the list of subordinate bodies not forming part of the Ministry of Economy’s structure (Decision No. 228, 30 Nov 2010).'),
    ('174','10 Jun 2026','Government Guarantees — SOCAR / Gran Tierra Energy Production Sharing Agreement',
     'Approves the Government’s guarantees and obligations under the PSA between SOCAR and Gran Tierra Energy (Azerbaijan) GmbH for exploration, development and production across parts of the Quba, Qusar, Shabran, Siyazan and Khizi districts.'),
    ('172','10 Jun 2026','Approval of the List of State Bodies with Designated Official Parking',
     'Approves the list of state bodies provided with official parking, the number of allocated spaces and the addresses of the relevant facilities.'),
    ('173','9 Jun 2026','Gas Supply — Implementing Amendments & Repeals (Law No. 233-VIIQ)',
     'Amends several CoM decisions to implement the Law on Gas Supply (No. 233-VIIQ, 8 Jul 2025) and repeals older rules (No. 87/1999 and Nos. 36 & 37/2023).'),
    ('171','9 Jun 2026','Rules on the Maintenance & Use of General-Use Railway Transport Infrastructure',
     'Adopts comprehensive rules governing the maintenance, operation and use of general-purpose railway infrastructure.'),
    ('170','9 Jun 2026','Amendment to the Mandatory Health Insurance Services Package',
     'Modifies the Mandatory Health Insurance Services Package approved by CoM Decision No. 5 (10 Jan 2020).'),
    ('169','9 Jun 2026','Implementation of the State Security Service Centralized Information & Digital Analytics System',
     'Amends CoM decisions to implement Presidential Decree No. 533 (21 Nov 2025) on the creation of the Centralized Information and Digital Analytics System of the State Security Service (DTX).'),
    ('168','9 Jun 2026','Regulation on the Energy Efficiency Information System',
     'Approves the Regulation establishing the framework for the State’s Energy Efficiency Information System.'),
]
subhead('Decisions (Qərarlar)')
for num,date,title,body in decisions:
    doc_item('Decision','2c3e50', f'№ {num}', date, title, body)

orders = [
    ('403s','10 Jun 2026','Execution of Presidential Decree No. 678 — ASAN / State Agency for Service to Citizens',
     'Gives effect to Presidential Decree No. 678 (2 Jun 2026), amending Decree No. 706 (5 Sep 2012) on the State Agency for Service to Citizens and Social Innovations (ASAN xidmət).'),
    ('402s','10 Jun 2026','Execution of Presidential Decree No. 675 — Finance Ministry Board Representation',
     'Implements Presidential Decree No. 675 (26 May 2026) on Ministry of Finance representation on the supervisory/oversight boards of state-owned and public legal entities.'),
    ('400s','9 Jun 2026','Execution of Presidential Order No. 1053 — Quba District Renaming (Law No. 404-VIIQ)',
     'Implements Presidential Order No. 1053 (1 Jun 2026) applying Law No. 404-VIIQ (12 May 2026) on renaming certain territorial units of the Quba district.'),
    ('389s','9 Jun 2026','Execution of Presidential Decree No. 672 — State Agency for Protection of Strategic Facilities',
     'Implements Presidential Decree No. 672 (25 May 2026), amending Decree No. 1510 (22 Dec 2021) on the statute, structure and staffing of the State Agency for the Protection of Strategic Facilities.'),
]
subhead('Orders (Sərəncamlar)')
for num,date,title,body in orders:
    doc_item('Order','34495e', f'№ {num}', date, title, body)

# ── MILLI MAJLIS ────────────────────────────────────────────────────────────
section('Parliament of Azerbaijan', 'Milli Majlis', 'Bills on the Agenda — 12 June 2026 Session')
p = para(after=4)
run(p, 'Bills placed before members at the 12 June 2026 session for discussion. No new laws or '
       'parliamentary decrees were enacted within the 8–13 June window (the most recent law, '
       'No. 415-VIIQD on Citizens’ Appeals, dates to 26 May and was brought into force by '
       'presidential decree on 10 June).', size=9, font='Arial', color=GREY)
bills = [
    '2025 State Budget Execution Report — annual accountability report on implementation of the 2025 state budget.',
    'Land Code Amendments — reform package on land lease, the land market, privatization and municipal land management. The week’s most commercially significant bill.',
    'Accession — International Convention on Harmful Anti-Fouling Systems on Ships (AFS).',
    'Establishment of the Azerbaijani Embassy in Lisbon, Portugal.',
    'Shushi Cultural Capital Status Amendment.',
    'Green Space Protection Law Amendment.',
    'Administrative Offenses Code Amendment.',
    'Public Service Law Amendment.',
    'Omnibus — Administrative Offenses, Information Protection & Child Safety Laws.',
    'Housing & Refugee Status Laws Amendment.',
]
for i,b in enumerate(bills,1):
    p = para(before=0, after=2)
    run(p, f'{i}.  ', size=9.5, bold=True)
    parts = b.split(' — ',1)
    run(p, parts[0], size=9.5, bold=True)
    if len(parts)>1:
        run(p, ' — '+parts[1], size=9.5)
    run(p, '   [Under Discussion]', size=7.5, italic=True, color=GREY, font='Arial')

# ── COURT + E-QANUN ─────────────────────────────────────────────────────────
section('Judicial / Official Database', 'Constitutional Court & e-Qanun', 'Source Status')
for title,bodytext in [
    ('⚠ Constitutional Court (constcourt.gov.az) — Unavailable.',
     'The decisions portal returned HTTP 503 (Service Unavailable) on repeated attempts via two methods '
     'at compilation (13 June 2026). No decisions from the reporting period could be retrieved. Consult the '
     'portal directly for any decisions issued 8–13 June 2026.'),
    ('⚠ e-Qanun (e-qanun.az) — Not Machine-Readable.',
     'The official legislative database serves an empty client-rendered (JavaScript) application shell with no '
     'static document listing. The full authoritative texts of all instruments above are nonetheless published '
     'there; rely on e-qanun.az for citation-grade text.'),
]:
    tbl = doc.add_table(rows=1, cols=1); cell = tbl.rows[0].cells[0]
    shade(cell, 'fff8e1'); cp = cell.paragraphs[0]
    run(cp, title+' ', size=8.5, bold=True, color=RGBColor(0x6b,0x59,0x00), font='Arial')
    run(cp, bodytext, size=8.5, color=RGBColor(0x6b,0x59,0x00), font='Arial')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── SPOTLIGHT ───────────────────────────────────────────────────────────────
section('Editorial Desk', 'Regulatory Spotlight', 'Key Themes, Week 24 / 2026')
spot = [
    ('1. Economic-Stimulus Trio (9 June)',
     'Three presidential decrees in two days target private-sector finance: enhanced credit guarantees and '
     'interest subsidies for entrepreneurs, new non-oil export incentives with transport-cost support, and a '
     'recapitalised agrarian-credit framework. Clients in export, agribusiness and SME finance should review '
     'eligibility once implementing rules are gazetted.'),
    ('2. Civil Aviation Governance',
     'The 12 June decree adopting Rules on the Management of Airports introduces a dedicated governance regime '
     'for airport operators. Parties to airport concession, ground-handling or infrastructure contracts should '
     'assess alignment with the new Rules and any transition provisions.'),
    ('3. Energy & Infrastructure Rulemaking',
     'The Cabinet adopted railway-infrastructure Rules (No. 171), restructured the gas-supply regime (No. 173), '
     'advanced the renewable-energy land bank in Hacıqabul/Jabrayil (Nos. 176–178), and confirmed Government '
     'guarantees for the SOCAR–Gran Tierra PSA (No. 174). A high-activity quarter for energy and projects teams.'),
    ('4. Capital Governance & Citizens’ Rights',
     'Two 10 June decrees reshape administration: a restructuring of the Baku City Executive Authority, and entry '
     'into force of the amended Law on Citizens’ Appeals with conforming changes to the 2016 appeals-handling '
     'Rules. Public-law and administrative-litigation practitioners should note revised procedural touchpoints.'),
]
for h,b in spot:
    p = para(before=6, after=2); run(p, h, size=10, bold=True)
    p2 = para(after=4); run(p2, b, size=9, color=RGBColor(0x44,0x44,0x44))

# ── FOOTER ──────────────────────────────────────────────────────────────────
hrule(para(before=10, after=2))
p = para(after=2)
run(p, 'Omni Law Gazette', size=8, bold=True, font='Arial', color=GREY)
run(p, ' is an internal digest compiled by Omni Law Firm for informational purposes only. It does not '
       'constitute legal advice. Descriptions are derived from official government portals and '
       'translated/summarised by the editorial desk; consult primary sources and the consolidated texts on '
       'e-qanun.az before relying on any item. Sources: president.az · nk.gov.az · meclis.gov.az · '
       'e-qanun.az · constcourt.gov.az', size=8, font='Arial', color=GREY)
p = para(after=0)
run(p, 'Issue No. 2  ·  Reporting period 8–13 June 2026  ·  Compiled 13 June 2026  ·  © 2026 Omni Law Firm',
    size=8, bold=True, font='Arial', color=GREY)

out = os.path.join(os.path.dirname(__file__), 'gazettes',
                   'Omni-Law-Gazette-Issue-2-2026-06-13.docx')
doc.save(out)
print('DOCX written:', out)
